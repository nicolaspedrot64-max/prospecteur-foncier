import gzip
import io
import json
import math
import re
import time
import unicodedata
import zipfile
from datetime import date
from pathlib import Path
from urllib.parse import urlparse

import pandas as pd
import pydeck as pdk
import requests
try:
    import duckdb
except ImportError:
    duckdb = None
import streamlit as st
from docx import Document
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    Flowable,
)
from pyproj import Geod
from shapely.geometry import shape, mapping
from shapely.ops import unary_union
from shapely.strtree import STRtree

APP_DIR = Path(__file__).parent
LETTER_TEMPLATE = APP_DIR / "lettre_sagec_modele.docx"

GEO_API = "https://geo.api.gouv.fr"
CADASTRE_BASE = "https://cadastre.data.gouv.fr/data/etalab-cadastre/latest/geojson/communes"
CADASTRE_BASE_FALLBACK = "https://files.data.gouv.fr/cadastre/etalab-cadastre/latest/geojson/communes"
GPU_API = "https://apicarto.ign.fr/api/gpu"
GEOCODAGE_API = "https://data.geopf.fr/geocodage"

# Zonage A/B/C officiel utilisé notamment pour le PTZ.
# Le dataset data.gouv.fr est maintenu par le ministère chargé du Logement.
PTZ_ZONAGE_DATASET_API = (
    "https://www.data.gouv.fr/api/1/datasets/"
    "liste-des-communes-selon-le-zonage-abc/"
)
# Fallback : liste nationale en vigueur au 26 juin 2026
# (arrêté du 23 juin 2026), au cas où la découverte dynamique du dataset échoue.
PTZ_ZONAGE_FALLBACK_RESOURCE = (
    "https://www.data.gouv.fr/api/1/datasets/r/"
    "13f7282b-8a25-43ab-9713-8bb4e476df55"
)

# Fichier open data DGFiP des parcelles détenues par des personnes morales.
# La ressource data.gouv.fr redirige vers le parquet courant.
FPMU_PARCELLES_RESOURCE = (
    "https://www.data.gouv.fr/api/1/datasets/r/"
    "913b0f65-3a22-4049-beb0-c33e5084df79"
)

GEOD = Geod(ellps="WGS84")

REGIONS = {
    "Nouvelle-Aquitaine": {
        "16": "Charente",
        "17": "Charente-Maritime",
        "19": "Corrèze",
        "23": "Creuse",
        "24": "Dordogne",
        "33": "Gironde",
        "40": "Landes",
        "47": "Lot-et-Garonne",
        "64": "Pyrénées-Atlantiques",
        "79": "Deux-Sèvres",
        "86": "Vienne",
        "87": "Haute-Vienne",
    },
    "Occitanie": {
        "09": "Ariège",
        "11": "Aude",
        "12": "Aveyron",
        "30": "Gard",
        "31": "Haute-Garonne",
        "32": "Gers",
        "34": "Hérault",
        "46": "Lot",
        "48": "Lozère",
        "65": "Hautes-Pyrénées",
        "66": "Pyrénées-Orientales",
        "81": "Tarn",
        "82": "Tarn-et-Garonne",
    },
}

st.set_page_config(
    page_title="Prospecteur Foncier V3.8",
    page_icon="🏗️",
    layout="wide",
)

# --------------------------
# Helpers réseau
# --------------------------
BDNB_API = "https://api.bdnb.io/v1/bdnb/donnees/batiment_groupe_complet"

HEADERS = {
    "User-Agent": "ProspecteurFoncier/3.8 (Streamlit; donnees publiques)",
    "Accept": "application/json, application/geo+json, */*",
}


def http_get_json(url, params=None, timeout=60):
    r = requests.get(url, params=params, headers=HEADERS, timeout=timeout)
    r.raise_for_status()
    return r.json()


@st.cache_data(ttl=24 * 3600, show_spinner=False)
def fetch_communes(dept_code):
    url = f"{GEO_API}/departements/{dept_code}/communes"
    data = http_get_json(
        url,
        params={
            "fields": "nom,code,codeEpci,codesPostaux,population",
            "format": "json",
        },
        timeout=30,
    )
    return sorted(data, key=lambda x: x.get("nom", ""))


@st.cache_data(ttl=24 * 3600, show_spinner=False)
def fetch_ptz_zonage():
    """
    Charge le zonage A/B/C officiel en vigueur.

    Le dataset est découvert dynamiquement via l'API data.gouv.fr afin que
    l'application puisse suivre les futures mises à jour. Si cette découverte
    échoue, on utilise la ressource nationale publiée pour le zonage en vigueur
    au 26 juin 2026.
    """
    resource_url = PTZ_ZONAGE_FALLBACK_RESOURCE
    source_label = "Zonage ABC en vigueur au 26 juin 2026"

    try:
        meta = http_get_json(PTZ_ZONAGE_DATASET_API, timeout=30)
        resources = meta.get("resources", []) or []

        csv_candidates = []
        for res in resources:
            fmt = str(res.get("format") or "").lower()
            title = str(res.get("title") or res.get("description") or "")
            url = str(res.get("url") or "")
            haystack = f"{title} {url}".lower()

            if fmt == "csv" or ".csv" in url.lower():
                score = 0
                if "ensemble" in haystack and "commune" in haystack:
                    score += 10
                if "zonage" in haystack and "abc" in haystack:
                    score += 5
                if "en vigueur" in haystack:
                    score += 5
                # Les ressources les plus récentes sont préférées en cas d'égalité.
                modified = str(
                    res.get("last_modified")
                    or res.get("modified")
                    or res.get("created_at")
                    or ""
                )
                csv_candidates.append((score, modified, url, title))

        if csv_candidates:
            csv_candidates.sort(key=lambda x: (x[0], x[1]), reverse=True)
            _, _, discovered_url, discovered_title = csv_candidates[0]
            if discovered_url:
                resource_url = discovered_url
                source_label = discovered_title or source_label
    except Exception:
        pass

    r = requests.get(resource_url, headers=HEADERS, timeout=45)
    r.raise_for_status()

    # Le CSV ministériel est séparé par des points-virgules.
    df = pd.read_csv(
        io.StringIO(r.text),
        sep=";",
        dtype=str,
        keep_default_na=False,
    )

    # Normaliser les noms de colonnes pour résister aux changements de millésime.
    cols = {str(c).strip().lower(): c for c in df.columns}

    code_col = None
    zone_col = None
    name_col = None

    for c in df.columns:
        lc = str(c).strip().lower()
        if code_col is None and (
            lc == "codgeo"
            or "code insee" in lc
            or ("code" in lc and "commune" in lc)
        ):
            code_col = c
        if zone_col is None and "zonage" in lc and "abc" in lc:
            zone_col = c
        if name_col is None and (
            lc == "libgeo"
            or "nom commune" in lc
            or ("libell" in lc and "commune" in lc)
        ):
            name_col = c

    if code_col is None or zone_col is None:
        raise RuntimeError(
            "Le fichier officiel de zonage PTZ a changé de structure "
            "et ses colonnes n'ont pas pu être identifiées."
        )

    result = {}
    for _, row in df.iterrows():
        insee = str(row.get(code_col) or "").strip().zfill(5)
        zone = str(row.get(zone_col) or "").strip()
        if not insee or not zone:
            continue

        # Uniformiser quelques variantes de typographie.
        z_norm = normalize_text(zone).replace(" ", "")
        if z_norm in {"abis", "a-bis"}:
            zone = "Abis"
        elif z_norm == "a":
            zone = "A"
        elif z_norm == "b1":
            zone = "B1"
        elif z_norm == "b2":
            zone = "B2"
        elif z_norm == "c":
            zone = "C"

        result[insee] = {
            "zone": zone,
            "nom": str(row.get(name_col) or "").strip() if name_col else "",
        }

    return result, source_label


def filter_communes_by_ptz(communes, ptz_choice, zonage_index):
    """
    Filtre une liste de communes geo.api.gouv.fr avec le zonage officiel ABC.

    Conformément au Code de la construction et de l'habitation, la zone A bis
    est incluse dans la zone A. Le filtre 'A' conserve donc A + A bis.
    """
    if ptz_choice == "Tous":
        return communes

    allowed = {"Abis", "A"} if ptz_choice == "A" else {ptz_choice}

    filtered = []
    for commune in communes:
        code = str(commune.get("code") or "").zfill(5)
        info = zonage_index.get(code, {})
        if info.get("zone") in allowed:
            c = dict(commune)
            c["zone_ptz"] = info.get("zone")
            filtered.append(c)

    return filtered


def _dept_from_insee(insee):
    insee = str(insee)
    if insee.startswith("97"):
        return insee[:3]
    return insee[:2]


@st.cache_data(ttl=6 * 3600, show_spinner=False)
def fetch_cadastre_layer(insee, layer):
    dep = _dept_from_insee(insee)
    filename = f"cadastre-{insee}-{layer}.json.gz"
    urls = [
        f"{CADASTRE_BASE}/{dep}/{insee}/{filename}",
        f"{CADASTRE_BASE_FALLBACK}/{dep}/{insee}/{filename}",
    ]
    last_error = None
    for url in urls:
        try:
            r = requests.get(url, headers=HEADERS, timeout=120)
            r.raise_for_status()
            return json.loads(gzip.decompress(r.content).decode("utf-8"))
        except Exception as exc:
            last_error = exc
    raise RuntimeError(
        f"Impossible de télécharger la couche cadastrale '{layer}' pour {insee}: {last_error}"
    )


def _normalise_parcelle_id(value):
    if value is None:
        return ""
    return re.sub(r"[^0-9A-Za-z]", "", str(value)).upper()


def _parse_parcelle_list(value):
    if value is None:
        return []
    if isinstance(value, list):
        return [_normalise_parcelle_id(v) for v in value if v]
    if isinstance(value, tuple):
        return [_normalise_parcelle_id(v) for v in value if v]

    txt = str(value).strip()
    if not txt:
        return []

    try:
        parsed = json.loads(txt)
        if isinstance(parsed, list):
            return [_normalise_parcelle_id(v) for v in parsed if v]
    except Exception:
        pass

    txt = txt.strip("{}[]")
    parts = re.split(r"[,;]", txt)
    return [_normalise_parcelle_id(p.strip().strip('"').strip("'")) for p in parts if p.strip()]


@st.cache_data(ttl=7 * 24 * 3600, show_spinner=False)
def fetch_bdnb_commune(insee):
    """
    Charge les usages principaux des bâtiments BDNB de la commune.
    Cette donnée sert uniquement à repérer les parcelles déjà occupées
    par du logement collectif.
    """
    fields = ",".join([
        "batiment_groupe_id",
        "code_commune_insee",
        "l_parcelle_id",
        "usage_principal_bdnb_open",
        "nb_log",
        "libelle_adr_principale_ban",
    ])

    rows = []
    limit = 1000
    offset = 0

    for _ in range(100):
        params = {
            "code_commune_insee": f"eq.{insee}",
            "select": fields,
            "limit": limit,
            "offset": offset,
            "order": "batiment_groupe_id.asc",
        }
        r = requests.get(BDNB_API, params=params, headers=HEADERS, timeout=120)
        r.raise_for_status()
        page = r.json()
        if not isinstance(page, list):
            raise RuntimeError("Réponse BDNB inattendue.")
        rows.extend(page)
        if len(page) < limit:
            break
        offset += limit

    return rows


def build_bdnb_parcel_index(rows):
    index = {}
    for row in rows or []:
        for pid in _parse_parcelle_list(row.get("l_parcelle_id")):
            if pid:
                index.setdefault(pid, []).append(row)
    return index


def detect_collective_housing(parcel_bdnb_rows):
    """
    Retourne True uniquement si la BDNB identifie du résidentiel collectif
    sur la parcelle. Tous les autres cas restent éligibles.
    """
    usages = []
    adresse = ""
    for row in parcel_bdnb_rows or []:
        usage = str(row.get("usage_principal_bdnb_open") or "").strip()
        if usage:
            usages.append(usage)
        if not adresse:
            adresse = str(row.get("libelle_adr_principale_ban") or "").strip()

    usages_norm = [normalize_text(u) for u in usages]
    collectif = any("residentiel collectif" in u for u in usages_norm)

    return {
        "collectif_existant": collectif,
        "usage_bdnb": " / ".join(sorted(set(usages))),
        "adresse_bdnb": adresse,
    }

def _sql_ident(name):
    return '"' + str(name).replace('"', '""') + '"'


def _pick_column(columns, exact=(), contains_all=(), contains_any=()):
    """
    Trouve une colonne dans un schéma qui peut évoluer légèrement entre millésimes.
    """
    lower = {str(c).lower(): str(c) for c in columns}
    for e in exact:
        if e.lower() in lower:
            return lower[e.lower()]

    for c in columns:
        lc = str(c).lower()
        if contains_all and all(x.lower() in lc for x in contains_all):
            return str(c)

    for c in columns:
        lc = str(c).lower()
        if contains_any and any(x.lower() in lc for x in contains_any):
            return str(c)

    return None


def _clean_owner_text(value):
    if value is None:
        return ""
    txt = str(value).strip()
    if txt.lower() in {"", "nan", "none", "null", "<na>"}:
        return ""
    return re.sub(r"\s+", " ", txt)


def _canonical_parcel_id(insee, prefixe, section, numero):
    insee = str(insee or "").strip()
    prefixe = re.sub(r"\D", "", str(prefixe or "")).zfill(3)[-3:]
    section = re.sub(r"[^0-9A-Za-z]", "", str(section or "")).upper().zfill(2)[-2:]
    numero = re.sub(r"\D", "", str(numero or "")).zfill(4)[-4:]
    return _normalise_parcelle_id(f"{insee}{prefixe}{section}{numero}")


@st.cache_data(ttl=24 * 3600, show_spinner=False)
def fetch_personnes_morales_commune(insee):
    """
    Interroge le parquet open data des parcelles appartenant à des personnes morales.
    Le fichier complet fait plusieurs centaines de Mo ; DuckDB ne télécharge que les
    blocs nécessaires grâce aux requêtes HTTP Range.
    """
    if duckdb is None:
        raise RuntimeError(
            "Le module duckdb n'est pas installé. Ajoute duckdb dans requirements.txt."
        )

    con = duckdb.connect(database=":memory:")
    try:
        # Les versions récentes de DuckDB chargent httpfs automatiquement ; on tente
        # explicitement le chargement puis on laisse l'auto-load prendre le relais.
        try:
            con.execute("LOAD httpfs")
        except Exception:
            try:
                con.execute("INSTALL httpfs")
                con.execute("LOAD httpfs")
            except Exception:
                pass

        parquet_sql = F"read_parquet('{FPMU_PARCELLES_RESOURCE}')"
        schema_df = con.execute(f"DESCRIBE SELECT * FROM {parquet_sql}").fetchdf()
        columns = schema_df["column_name"].astype(str).tolist()

        col_insee = _pick_column(
            columns,
            exact=("code_insee", "code_commune_insee", "codcomm"),
            contains_all=("insee",),
        )
        col_parcelle = _pick_column(
            columns,
            exact=("parcelle_id", "id_parcelle", "parcelle", "idpar"),
            contains_all=("parcelle", "id"),
        )
        col_prefixe = _pick_column(
            columns,
            exact=("prefixe", "prefixe_section", "ccopre"),
            contains_any=("prefix", "ccopre"),
        )
        col_section = _pick_column(
            columns,
            exact=("section", "ccosec"),
            contains_any=("section", "ccosec"),
        )
        col_numero = _pick_column(
            columns,
            exact=("numero", "numero_parcelle", "dnupla"),
            contains_any=("numero_parcelle", "dnupla"),
        )
        col_denom = _pick_column(
            columns,
            exact=("denomination", "denomination_proprietaire", "dforme"),
            contains_any=("denomination", "raison_sociale", "proprietaire_nom"),
        )
        col_forme = _pick_column(
            columns,
            exact=("forme_juridique", "forme_juridique_abregee", "groupe_personne"),
            contains_any=("forme_juridique", "forme juridique", "groupe_personne"),
        )
        col_siren = _pick_column(
            columns,
            exact=("siren", "numero_siren"),
            contains_any=("siren",),
        )

        if not col_denom:
            raise RuntimeError(
                "Le fichier personnes morales est accessible mais la colonne de dénomination "
                "n'a pas pu être identifiée automatiquement."
            )

        wanted = []
        aliases = {}
        for alias, col in [
            ("code_insee", col_insee),
            ("parcelle_id", col_parcelle),
            ("prefixe", col_prefixe),
            ("section", col_section),
            ("numero", col_numero),
            ("denomination", col_denom),
            ("forme_juridique", col_forme),
            ("siren", col_siren),
        ]:
            if col:
                wanted.append(f"{_sql_ident(col)} AS {_sql_ident(alias)}")
                aliases[alias] = col

        if not wanted:
            return []

        where = ""
        params = []

        if col_insee:
            where = f"WHERE CAST({_sql_ident(col_insee)} AS VARCHAR) = ?"
            params = [str(insee)]
        elif col_parcelle:
            # Le début de l'identifiant cadastral contient le code commune INSEE
            # dans la représentation utilisée par les jeux cadastraux ouverts.
            where = f"WHERE CAST({_sql_ident(col_parcelle)} AS VARCHAR) LIKE ?"
            params = [f"{insee}%"]

        query = f"SELECT {', '.join(wanted)} FROM {parquet_sql} {where}"
        df = con.execute(query, params).fetchdf()
        return df.to_dict("records")
    finally:
        con.close()


def build_personnes_morales_index(insee, rows):
    """
    Indexe les propriétaires personnes morales par identifiant cadastral normalisé.
    Plusieurs propriétaires sont concaténés lorsqu'ils existent.
    """
    temp = {}

    for r in rows or []:
        pid = _normalise_parcelle_id(r.get("parcelle_id"))
        if not pid:
            pid = _canonical_parcel_id(
                insee,
                r.get("prefixe"),
                r.get("section"),
                r.get("numero"),
            )
        if not pid:
            continue

        denom = _clean_owner_text(r.get("denomination"))
        forme = _clean_owner_text(r.get("forme_juridique"))
        siren = _clean_owner_text(r.get("siren"))

        if not denom:
            continue

        key = (denom, forme, siren)
        temp.setdefault(pid, set()).add(key)

    index = {}
    for pid, owners in temp.items():
        owners = sorted(owners)
        names = [o[0] for o in owners if o[0]]
        formes = [o[1] for o in owners if o[1]]
        sirens = [o[2] for o in owners if o[2]]

        joined_name = " / ".join(dict.fromkeys(names))
        joined_forme = " / ".join(dict.fromkeys(formes))
        joined_siren = " / ".join(dict.fromkeys(sirens))

        upper = normalize_text(joined_name + " " + joined_forme)
        is_commune = any(
            token in upper
            for token in [
                "commune de ",
                "commune d ",
                "mairie",
                "ville de ",
                "commune",
            ]
        )

        index[pid] = {
            "proprietaire_personne_morale": joined_name,
            "forme_juridique_proprietaire": joined_forme,
            "siren_proprietaire": joined_siren,
            "proprietaire_commune": is_commune,
            "proprietaire_type": "Commune / collectivité" if is_commune else "Société / personne morale",
        }

    return index


def enrich_with_personnes_morales(df, insee, owner_index):
    if df is None:
        return df

    df = df.copy()
    defaults = {
        "proprietaire_personne_morale": "",
        "forme_juridique_proprietaire": "",
        "siren_proprietaire": "",
        "proprietaire_type": "",
        "proprietaire_commune": False,
    }
    for c, default in defaults.items():
        if c not in df.columns:
            df[c] = default

    if df.empty:
        return df

    for idx, row in df.iterrows():
        pid = _normalise_parcelle_id(row.get("reference"))
        info = owner_index.get(pid)
        if not info:
            # Fallback : reconstruire un identifiant à partir des champs cadastraux.
            pid2 = _canonical_parcel_id(
                insee,
                "",
                row.get("section"),
                row.get("numero"),
            )
            info = owner_index.get(pid2)

        if info:
            for k, v in info.items():
                df.at[idx, k] = v

    return df

def gpu_get(layer, params, timeout=90):
    """
    Appel robuste à l'API Carto GPU.
    Les erreurs 500/502/503/504 sont parfois temporaires : on réessaie avant abandon.
    """
    url = f"{GPU_API}/{layer}"
    last_error = None

    # GET est la voie principale documentée.
    for attempt in range(3):
        try:
            r = requests.get(url, params=params, headers=HEADERS, timeout=timeout)
            if r.ok:
                return r.json()

            last_error = requests.HTTPError(
                f"{r.status_code} {r.reason} pour {r.url}",
                response=r,
            )
            if r.status_code not in {429, 500, 502, 503, 504}:
                break
        except Exception as exc:
            last_error = exc

        time.sleep(0.7 * (attempt + 1))

    # Certaines requêtes géométriques peuvent fonctionner en POST lorsque GET échoue.
    try:
        r2 = requests.post(url, json=params, headers=HEADERS, timeout=timeout)
        if r2.ok:
            return r2.json()
        last_error = requests.HTTPError(
            f"{r2.status_code} {r2.reason} pour {r2.url}",
            response=r2,
        )
    except Exception as exc:
        last_error = exc

    if isinstance(last_error, Exception):
        raise last_error
    raise RuntimeError(f"Échec de l'appel GPU {layer}")


@st.cache_data(ttl=6 * 3600, show_spinner=False)
def fetch_gpu_document(insee, commune_geojson, code_epci=None):
    """
    Résout le document d'urbanisme en vigueur.

    1. Essaie l'intersection ponctuelle /gpu/document?geom=...
    2. Si ce service renvoie une erreur (ex. HTTP 500), teste directement les
       partitions normalisées GPU :
       - DU_<code EPCI> pour un PLUi ;
       - DU_<code INSEE> pour un PLU communal.

    Le second chemin évite qu'une panne du seul endpoint /document bloque toute l'application.
    """
    errors = []

    # 1) Recherche géométrique normale.
    features = commune_geojson.get("features", [])
    if features:
        try:
            commune_geom = shape(features[0]["geometry"])
            point = commune_geom.representative_point()
            geom = json.dumps(mapping(point), separators=(",", ":"))

            data = gpu_get("document", {"geom": geom})
            docs = data.get("features", []) or []

            local_docs = []
            for f in docs:
                p = f.get("properties", {}) or {}
                typedoc = str(p.get("typedoc") or p.get("type_doc") or "").upper()
                if "SCOT" in typedoc:
                    continue
                local_docs.append(f)

            candidates = local_docs or docs
            if candidates:
                def doc_sort_key(f):
                    p = f.get("properties", {}) or {}
                    d = (
                        p.get("datvalid")
                        or p.get("datappro")
                        or p.get("date_appro")
                        or p.get("datefin")
                        or ""
                    )
                    return str(d)

                candidates = sorted(candidates, key=doc_sort_key, reverse=True)
                props = candidates[0].get("properties", {}) or {}
                partition = props.get("partition")
                if partition:
                    return partition, candidates
        except Exception as exc:
            errors.append(f"recherche géométrique: {exc}")

    # 2) Fallback par partitions. Pour les PLUi, le GPU utilise DU_<SIREN EPCI>.
    partition_candidates = []
    if code_epci:
        partition_candidates.append(f"DU_{code_epci}")
    partition_candidates.append(f"DU_{insee}")

    seen = set()
    for partition in partition_candidates:
        if not partition or partition in seen:
            continue
        seen.add(partition)

        try:
            # On filtre déjà sur la commune pour éviter de charger tout le PLUi.
            probe = gpu_get(
                "zone-urba",
                {"partition": partition, "insee": str(insee)},
                timeout=90,
            )
            feats = probe.get("features", []) or []
            if feats:
                return partition, []
        except Exception as exc:
            errors.append(f"{partition}: {exc}")

    details = " | ".join(errors[-4:]) if errors else "aucun zonage trouvé"
    raise RuntimeError(
        "Impossible de résoudre le document PLU/PLUi. "
        f"Partitions testées : {', '.join(partition_candidates)}. Détail : {details}"
    )


@st.cache_data(ttl=6 * 3600, show_spinner=False)
def fetch_gpu_zones(partition, insee=None):
    # API GPU : limite haute de plusieurs milliers d'objets.
    all_features = []
    start = 0
    while True:
        params = {"partition": partition}
        if insee:
            params["insee"] = str(insee)
        if start:
            params["_start"] = start

        data = gpu_get("zone-urba", params)
        feats = data.get("features", []) or []
        all_features.extend(feats)

        returned = data.get("numberReturned", len(feats))
        total = data.get("totalFeatures", data.get("numberMatched"))

        if not feats or returned == 0:
            break
        if total is not None and len(all_features) >= int(total):
            break
        if len(feats) < 1000:
            break

        start += len(feats)
        if start >= 5000:
            break

    return {"type": "FeatureCollection", "features": all_features}


@st.cache_data(ttl=6 * 3600, show_spinner=False)
def fetch_gpu_layer(layer, partition, insee=None):
    """
    Charge une couche GPU par partition et commune.
    Utilisé notamment pour prescription-surf / lin / pct.
    """
    all_features = []
    start = 0
    while True:
        params = {"partition": partition}
        if insee:
            params["insee"] = str(insee)
        if start:
            params["_start"] = start

        data = gpu_get(layer, params)
        feats = data.get("features", []) or []
        all_features.extend(feats)

        returned = data.get("numberReturned", len(feats))
        total = data.get("totalFeatures", data.get("numberMatched"))

        if not feats or returned == 0:
            break
        if total is not None and len(all_features) >= int(total):
            break
        if len(feats) < 1000:
            break

        start += len(feats)
        if start >= 10000:
            break

    return {"type": "FeatureCollection", "features": all_features}


@st.cache_data(ttl=6 * 3600, show_spinner=False)
def fetch_gpu_graphic_prescriptions(partition, insee=None):
    """
    Les codes CNIG recherchés sont :
    - 38 / 02 : emprise au sol maximale
    - 39 / 02 : hauteur maximale
    Les occurrences 39-50 / 39-51 sont aussi conservées lorsqu'elles existent
    dans certains standards/documents (façade / faîtage ou construction).
    """
    result = {}
    for layer in ["prescription-surf", "prescription-lin", "prescription-pct"]:
        try:
            result[layer] = fetch_gpu_layer(layer, partition, insee)
        except Exception:
            result[layer] = {"type": "FeatureCollection", "features": []}
    return result


@st.cache_data(ttl=6 * 3600, show_spinner=False)
def fetch_gpu_municipality(insee):
    try:
        return gpu_get("municipality", {"insee": insee})
    except Exception:
        return {"type": "FeatureCollection", "features": []}


@st.cache_data(ttl=30 * 24 * 3600, show_spinner=False)
def reverse_geocode(lat, lon, citycode=None):
    params = {
        "lat": lat,
        "lon": lon,
        "index": "address",
        "limit": 1,
    }
    if citycode:
        params["citycode"] = citycode
    data = http_get_json(f"{GEOCODAGE_API}/reverse", params=params, timeout=30)
    feats = data.get("features", []) or []
    if not feats:
        return ""
    p = feats[0].get("properties", {}) or {}
    return (
        p.get("label")
        or p.get("name")
        or " ".join(
            str(x)
            for x in [p.get("housenumber"), p.get("street"), p.get("postcode"), p.get("city")]
            if x
        )
    )


# --------------------------
# Helpers géographiques
# --------------------------
def _clean_psc_code(value):
    txt = re.sub(r"\D", "", str(value or ""))
    if not txt:
        return ""
    return txt.zfill(2)[-2:]


def _parse_optional_psc_attrs(props):
    """
    Les attributs supplémentaires CNIG sont stockés par paires
    LIB_ATTRn / LIB_VALn lorsqu'ils sont exposés par la source.
    """
    lower = {str(k).lower(): v for k, v in (props or {}).items()}
    pairs = []
    for i in range(1, 21):
        a = lower.get(f"lib_attr{i}")
        v = lower.get(f"lib_val{i}")
        if a not in [None, ""] or v not in [None, ""]:
            pairs.append((str(a or ""), str(v or "")))
    return pairs


def _parse_percent_value(text):
    txt = normalize_text(text)
    m = re.search(r"(?<!\d)(\d{1,3}(?:[.,]\d+)?)\s*%", txt)
    if m:
        val = float(m.group(1).replace(",", "."))
        if 0 < val <= 100:
            return val

    # coefficient décimal : 0,4 => 40 %
    m = re.search(r"(?:ces|coef(?:ficient)?[^0-9]{0,20})(0[.,]\d+)", txt)
    if m:
        val = float(m.group(1).replace(",", ".")) * 100.0
        if 0 < val <= 100:
            return val
    return None


def _parse_rplus_levels(text):
    txt = normalize_text(text)
    m = re.search(r"\br\s*\+\s*(\d{1,2})\b", txt)
    if m:
        return int(m.group(1)) + 1
    if re.search(r"\brdc\b|\brez[- ]de[- ]chaussee\b", txt):
        # Ne conclure à R seul que si le texte parle explicitement de hauteur/niveaux.
        if "niveau" in txt or "hauteur" in txt:
            return 1
    m = re.search(r"(?<!\d)(\d{1,2})\s+niveaux?\b", txt)
    if m:
        n = int(m.group(1))
        if 1 <= n <= 20:
            return n
    return None


def _parse_height_m(text):
    txt = normalize_text(text)
    # Chercher une valeur liée à la hauteur, pas n'importe quelle cote du libellé.
    patterns = [
        r"hauteur.{0,80}?(\d{1,2}(?:[.,]\d+)?)\s*m\b",
        r"(\d{1,2}(?:[.,]\d+)?)\s*m\b.{0,60}?hauteur",
        r"h\s*[=:]\s*(\d{1,2}(?:[.,]\d+)?)\s*m\b",
    ]
    for pat in patterns:
        m = re.search(pat, txt)
        if m:
            val = float(m.group(1).replace(",", "."))
            if 2 <= val <= 80:
                return val
    return None


def _rule_from_prescription(feature, layer_name):
    props = feature.get("properties", {}) or {}
    typepsc = _clean_psc_code(props.get("typepsc"))
    stypepsc = _clean_psc_code(props.get("stypepsc"))

    if typepsc not in {"38", "39"}:
        return None

    geom = valid_shape(feature)
    if geom is None:
        return None

    libelle = str(props.get("libelle") or "")
    txt = str(props.get("txt") or "")
    nature = str(props.get("nature") or "")
    urlfic = str(props.get("urlfic") or "")
    base_text = " | ".join(x for x in [libelle, txt, nature] if x)

    optionals = _parse_optional_psc_attrs(props)
    optional_text = " | ".join(f"{a}={v}" for a, v in optionals)
    all_text = " | ".join(x for x in [base_text, optional_text] if x)

    rule = {
        "geom": geom,
        "layer": layer_name,
        "typepsc": typepsc,
        "stypepsc": stypepsc,
        "libelle": libelle,
        "txt": txt,
        "urlfic": urlfic,
        "raw": all_text,
        "emprise_pct": None,
        "height_m": None,
        "levels": None,
        "confidence": 0,
        "method": "",
    }

    if typepsc == "38":
        # 38-02 = emprise maximale. 38-00 générique accepté si valeur explicite.
        if stypepsc not in {"00", "02"}:
            return None

        for attr, val in optionals:
            na = normalize_text(attr).replace(" ", "_")
            if "empris" in na and ("max" in na or stypepsc == "02"):
                pct = _parse_percent_value(val)
                if pct is not None:
                    rule["emprise_pct"] = pct
                    rule["confidence"] = 100
                    rule["method"] = f"Attribut graphique {attr}"
                    return rule

        pct = _parse_percent_value(all_text)
        if pct is not None:
            rule["emprise_pct"] = pct
            rule["confidence"] = 88 if stypepsc == "02" else 75
            rule["method"] = "Étiquette/libellé prescription graphique"
            return rule
        return rule

    # type 39 : hauteur.
    # On privilégie 39-02 (max). 50/51 sont tolérés pour documents qui les publient.
    if stypepsc not in {"00", "02", "50", "51"}:
        return None

    for attr, val in optionals:
        na = normalize_text(attr).replace(" ", "_")
        if "hauteur" in na and ("rpl" in na or "etage" in na):
            levels = _parse_rplus_levels(val)
            if levels is not None:
                rule["levels"] = levels
                rule["confidence"] = 100
                rule["method"] = f"Attribut graphique {attr}"
        if "hauteur" in na and ("metre" in na or "meter" in na or "max" in na):
            h = _parse_height_m(f"hauteur {val}")
            if h is not None:
                rule["height_m"] = h
                rule["confidence"] = max(rule["confidence"], 100)
                rule["method"] = f"Attribut graphique {attr}"

    if rule["levels"] is None:
        levels = _parse_rplus_levels(all_text)
        if levels is not None:
            rule["levels"] = levels
            rule["confidence"] = max(rule["confidence"], 90 if stypepsc == "02" else 80)
            rule["method"] = "Étiquette/libellé prescription graphique"

    if rule["height_m"] is None:
        h = _parse_height_m(all_text)
        if h is not None:
            rule["height_m"] = h
            rule["confidence"] = max(rule["confidence"], 88 if stypepsc == "02" else 75)
            if not rule["method"]:
                rule["method"] = "Étiquette/libellé prescription graphique"

    return rule


def build_graphic_rule_indexes(prescriptions):
    """
    Prépare un index spatial. Les règles surfaciques sont la source principale
    pour un calcul par parcelle. Les règles linéaires/ponctuelles de hauteur
    sont conservées comme indications supplémentaires.
    """
    surf_rules = []
    other_height_rules = []

    for layer_name, fc in (prescriptions or {}).items():
        for feat in fc.get("features", []) or []:
            rule = _rule_from_prescription(feat, layer_name)
            if not rule:
                continue
            if layer_name == "prescription-surf":
                surf_rules.append(rule)
            elif rule["typepsc"] == "39":
                other_height_rules.append(rule)

    surf_geoms = [r["geom"] for r in surf_rules]
    other_geoms = [r["geom"] for r in other_height_rules]

    return {
        "surf_rules": surf_rules,
        "surf_tree": STRtree(surf_geoms) if surf_geoms else None,
        "other_height_rules": other_height_rules,
        "other_tree": STRtree(other_geoms) if other_geoms else None,
    }


def _select_surface_rule(parcel_geom, tree, rules, kind):
    if tree is None:
        return None

    idxs = tree.query(parcel_geom, predicate="intersects")
    candidates = []
    parcel_area = max(geodesic_area_m2(parcel_geom), 0.01)

    for idx in idxs:
        rule = rules[int(idx)]
        if kind == "emprise" and rule.get("emprise_pct") is None:
            continue
        if kind == "height" and rule.get("levels") is None and rule.get("height_m") is None:
            continue

        try:
            inter = parcel_geom.intersection(rule["geom"])
            overlap_area = geodesic_area_m2(inter)
        except Exception:
            overlap_area = 0.0

        if overlap_area <= 0:
            continue

        ratio = overlap_area / parcel_area
        candidates.append((ratio, rule))

    if not candidates:
        return None

    # Priorité à la prescription couvrant la plus grande part de la parcelle.
    candidates.sort(key=lambda x: (-x[0], -x[1].get("confidence", 0)))
    ratio, best = candidates[0]

    # Une prescription graphique sur une toute petite portion ne doit pas être
    # appliquée à toute la parcelle. On la signale seulement si < 50 %.
    result = dict(best)
    result["coverage_ratio"] = ratio
    result["partial"] = ratio < 0.50
    return result


def _select_line_point_height(parcel_geom, tree, rules):
    if tree is None:
        return None
    idxs = tree.query(parcel_geom, predicate="intersects")
    candidates = []
    for idx in idxs:
        rule = rules[int(idx)]
        if rule.get("levels") is None and rule.get("height_m") is None:
            continue
        candidates.append(rule)
    if not candidates:
        return None

    # En cas de plusieurs plafonds, prendre le plus restrictif.
    def key(r):
        if r.get("levels") is not None:
            return (0, r["levels"])
        return (1, r.get("height_m") or 999)
    return sorted(candidates, key=key)[0]


def enrich_with_graphic_plu_rules(
    df,
    feature_map,
    rule_indexes,
    ratio_sdp_pct,
    ratio_shab_pct,
    shab_par_logement,
    floor_height_m,
):
    df = df.copy()

    cols = {
        "emprise_plu_pct": None,
        "emprise_graphique_couverture_pct": None,
        "hauteur_plu_m": None,
        "niveaux_plu": None,
        "surface_brute_plu_m2": None,
        "sdp_plu_m2": None,
        "shab_plu_m2": None,
        "logements_plu": None,
        "gabarit_plu_statut": "Non graphisé / à vérifier",
        "gabarit_plu_source": "",
        "gabarit_plu_confiance": 0,
        "prescription_emprise": "",
        "prescription_hauteur": "",
        "prescription_url": "",
    }
    for c, default in cols.items():
        if c not in df.columns:
            df[c] = default

    if df.empty:
        return df

    surf_tree = rule_indexes.get("surf_tree")
    surf_rules = rule_indexes.get("surf_rules", [])
    other_tree = rule_indexes.get("other_tree")
    other_rules = rule_indexes.get("other_height_rules", [])

    for idx, row in df.iterrows():
        feat = feature_map.get(row["reference"])
        if not feat:
            continue
        pg = valid_shape(feat)
        if pg is None:
            continue

        erule = _select_surface_rule(pg, surf_tree, surf_rules, "emprise")
        hrule = _select_surface_rule(pg, surf_tree, surf_rules, "height")

        # Si aucune hauteur surfacique, regarder les prescriptions linéaires/ponctuelles.
        if hrule is None:
            hrule = _select_line_point_height(pg, other_tree, other_rules)
            if hrule:
                hrule = dict(hrule)
                hrule["coverage_ratio"] = None
                hrule["partial"] = False

        if erule:
            df.at[idx, "emprise_plu_pct"] = erule.get("emprise_pct")
            df.at[idx, "emprise_graphique_couverture_pct"] = round(
                100 * erule.get("coverage_ratio", 0), 1
            )
            df.at[idx, "prescription_emprise"] = (
                f"{erule.get('typepsc')}-{erule.get('stypepsc')} "
                f"{erule.get('libelle') or erule.get('txt') or ''}"
            ).strip()
            if erule.get("urlfic"):
                df.at[idx, "prescription_url"] = erule.get("urlfic")

        if hrule:
            levels = hrule.get("levels")
            height = hrule.get("height_m")
            if levels is None and height is not None:
                levels = max(1, int(math.floor(float(height) / max(float(floor_height_m), 0.1))))
            df.at[idx, "hauteur_plu_m"] = height
            df.at[idx, "niveaux_plu"] = levels
            df.at[idx, "prescription_hauteur"] = (
                f"{hrule.get('typepsc')}-{hrule.get('stypepsc')} "
                f"{hrule.get('libelle') or hrule.get('txt') or ''}"
            ).strip()
            if not df.at[idx, "prescription_url"] and hrule.get("urlfic"):
                df.at[idx, "prescription_url"] = hrule.get("urlfic")

        # Ne pas généraliser une prescription surfacique qui ne couvre qu'un petit morceau.
        partial_emprise = bool(erule and erule.get("partial"))
        partial_height = bool(hrule and hrule.get("partial"))

        emprise = df.at[idx, "emprise_plu_pct"]
        levels = df.at[idx, "niveaux_plu"]

        if (
            emprise not in [None, ""]
            and levels not in [None, ""]
            and not partial_emprise
            and not partial_height
        ):
            terrain = float(row["surface_m2"])
            footprint = terrain * float(emprise) / 100.0
            gross = footprint * float(levels)
            sdp = gross * float(ratio_sdp_pct) / 100.0
            shab = sdp * float(ratio_shab_pct) / 100.0
            logements = math.floor(shab / max(float(shab_par_logement), 1.0))

            df.at[idx, "surface_brute_plu_m2"] = round(gross)
            df.at[idx, "sdp_plu_m2"] = round(sdp)
            df.at[idx, "shab_plu_m2"] = round(shab)
            df.at[idx, "logements_plu"] = int(max(0, logements))
            df.at[idx, "gabarit_plu_statut"] = "Calcul graphique"
            df.at[idx, "gabarit_plu_source"] = "Prescriptions CNIG 38/39"
            confs = [
                x.get("confidence", 0)
                for x in [erule, hrule]
                if x is not None
            ]
            df.at[idx, "gabarit_plu_confiance"] = min(confs) if confs else 0
        else:
            reasons = []
            if emprise in [None, ""]:
                reasons.append("emprise 38-02 absente/non chiffrée")
            if levels in [None, ""]:
                reasons.append("hauteur 39-02 absente/non chiffrée")
            if partial_emprise:
                reasons.append("emprise graphique partielle")
            if partial_height:
                reasons.append("hauteur graphique partielle")
            df.at[idx, "gabarit_plu_statut"] = "À vérifier — " + ", ".join(reasons)

    return df

def geodesic_area_m2(geom):
    try:
        area, _ = GEOD.geometry_area_perimeter(geom)
        return abs(float(area))
    except Exception:
        return 0.0


def valid_shape(feature):
    try:
        g = shape(feature["geometry"])
        if g.is_empty:
            return None
        if not g.is_valid:
            g = g.buffer(0)
        return g if not g.is_empty else None
    except Exception:
        return None


def normalize_text(s):
    s = str(s or "").lower()
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", s).strip()


NON_RESIDENTIAL_KEYWORDS = [
    "activite",
    "activites",
    "industri",
    "artisan",
    "econom",
    "commercial",
    "commerce",
    "logistique",
    "portuaire",
    "aeroport",
    "ferroviaire",
    "carriere",
    "dechet",
    "camping",
    "equipement",
    "sport",
    "tourisme",
    "loisir",
    "agricol",
    "naturel",
]

HABITAT_KEYWORDS = [
    "habitat",
    "habitation",
    "residentiel",
    "residentielle",
    "resident",
    "logement",
    "logements",
    "mixte habitat",
    "mixte resident",
]


ECONOMIC_VOCATION_KEYWORDS = [
    "zone d activite",
    "zone d activites",
    "secteur d activite",
    "secteur d activites",
    "activite economique",
    "activites economiques",
    "vocation economique",
    "zone economique",
    "parc d activite",
    "parc d activites",
    "zone industrielle",
    "secteur industriel",
    "zone artisanale",
    "secteur artisanal",
    "zone commerciale",
    "secteur commercial",
    "zone logistique",
    "secteur logistique",
    "plateforme logistique",
    "zone tertiaire",
    "secteur tertiaire",
    "pole economique",
    "pole d activite",
    "pole d activites",
]


def detect_economic_vocation(props):
    """
    Exclusion forte des secteurs dont la vocation dominante est l'activité économique.

    Sources structurées :
    - ancien standard CNIG : DESTDOMI=02 = activité ;
    - standard CNIG récent : FORMDOMI 0200 à 0203 = activité /
      industrielle-logistique-commerciale / commerces / bureaux.

    En complément, on analyse le libellé et le nom long du zonage.
    Les secteurs mixtes habitat/activité (DESTDOMI=03 / FORMDOMI=0300)
    ne sont pas considérés comme purement économiques par cette fonction.
    """
    destdomi = str(props.get("destdomi") or "").strip()
    formdomi = str(props.get("formdomi") or "").strip()
    libelle = str(props.get("libelle") or "")
    libelong = str(props.get("libelong") or "")
    text = normalize_text(" ".join([libelle, libelong]))

    dd = re.sub(r"\D", "", destdomi)
    fd = re.sub(r"\D", "", formdomi)

    if dd in {"2", "02"}:
        return True, "DESTDOMI=02 : vocation dominante activité"

    if fd.startswith("02") and len(fd) >= 2:
        labels = {
            "0200": "activité",
            "0201": "activité industrielle / logistique / commerciale",
            "0202": "activité commerces",
            "0203": "activité bureaux",
        }
        return True, f"FORMDOMI={formdomi} : {labels.get(fd, 'activité économique')}"

    for keyword in ECONOMIC_VOCATION_KEYWORDS:
        if keyword in text:
            return True, f"Libellé PLU : {libelle or libelong}"

    return False, ""

# Standard CNIG 2024/2025 :
# 20 = destination Habitation ; 21 = sous-destination Logement ; 22 = Hébergement.
DESTINATION_CODES = [
    "10", "11", "12",
    "20", "21", "22",
    "30", "31", "32", "33", "34", "35", "36", "37",
    "40", "41", "42", "43", "44", "45", "46", "47",
    "50", "51", "52", "53", "54", "55",
    "99",
]


def parse_destination_codes(value):
    """
    Extrait les codes CNIG de destination/sous-destination.
    Fonctionne avec des listes séparées ou des chaînes compactées.
    """
    if value is None:
        return set()

    if isinstance(value, (list, tuple, set)):
        txt = " ".join(str(x) for x in value)
    else:
        txt = str(value)

    found = set()
    i = 0
    while i < len(txt):
        matched = False
        for c in DESTINATION_CODES:
            if txt.startswith(c, i):
                found.add(c)
                i += len(c)
                matched = True
                break
        if not matched:
            i += 1
    return found


def _normalise_typezone(typezone):
    t = str(typezone or "").strip()
    if not t:
        return ""
    # Conserver la casse utile pour AUc/AUs
    tu = t.upper()
    if tu == "U":
        return "U"
    if tu == "AUC":
        return "AUc"
    if tu == "AUS":
        return "AUs"
    if tu == "AU":
        # Données plus anciennes : AU générique, traité comme zone à urbaniser conditionnelle.
        return "AU"
    if tu == "A":
        return "A"
    if tu == "N":
        return "N"
    return t


def analyse_habitat_zone(props, include_au=True, include_conditionnel=True):
    """
    Moteur Habitat à haute précision.

    Ordre de preuve :
    1) données structurées CNIG récentes DESTOUI / DESTCDT / DESTNON ;
    2) ancien attribut DESTDOMI (01 habitat, 03 mixte habitat/activité) ;
    3) libellé / nom long de la zone ;
    4) si aucune preuve explicite : zone classée "indéterminée", donc rejetée en mode strict.
    """
    raw_typezone = props.get("typezone")
    typezone = _normalise_typezone(raw_typezone)

    libelle = str(props.get("libelle") or "")
    libelong = str(props.get("libelong") or "")
    destdomi = str(props.get("destdomi") or "").strip()
    destoui_raw = props.get("destoui")
    destcdt_raw = props.get("destcdt")
    destnon_raw = props.get("destnon")

    text = normalize_text(" ".join([libelle, libelong]))

    destoui = parse_destination_codes(destoui_raw)
    destcdt = parse_destination_codes(destcdt_raw)
    destnon = parse_destination_codes(destnon_raw)

    # Priorité absolue : éliminer les secteurs dont la vocation dominante
    # est l'activité économique, même si certaines destinations y sont ponctuellement admises.
    economic_vocation, economic_reason = detect_economic_vocation(props)
    if economic_vocation:
        return {
            "habitat_eligible": False,
            "habitat_statut": "Exclue — secteur à vocation d'activités économiques",
            "habitat_preuve": economic_reason,
            "habitat_confiance": 100,
            "typezone_normalise": typezone,
            "economic_vocation": True,
        }

    # 0. Type de zone : on exclut les zones non immédiatement destinées à bâtir du logement.
    if typezone in {"A", "N"}:
        return {
            "habitat_eligible": False,
            "habitat_statut": "Exclue — zone agricole/naturelle",
            "habitat_preuve": f"TYPEZONE={typezone}",
            "habitat_confiance": 100,
            "typezone_normalise": typezone,
            "economic_vocation": False,
        }

    if typezone == "AUs":
        return {
            "habitat_eligible": False,
            "habitat_statut": "Exclue — zone AU bloquée",
            "habitat_preuve": "TYPEZONE=AUs : ouverture subordonnée à modification/révision du PLU",
            "habitat_confiance": 100,
            "typezone_normalise": typezone,
            "economic_vocation": False,
        }

    if typezone in {"AUc", "AU"} and not include_au:
        return {
            "habitat_eligible": False,
            "habitat_statut": "Exclue — zone AU désactivée dans les critères",
            "habitat_preuve": f"TYPEZONE={typezone}",
            "habitat_confiance": 100,
            "typezone_normalise": typezone,
            "economic_vocation": False,
        }

    # 1. Standards CNIG récents : destination Habitation / Logement.
    housing_codes = {"20", "21"}

    # Interdiction explicite prioritaire lorsqu'aucune autorisation explicite ne la contredit.
    if (destnon & housing_codes) and not (destoui & housing_codes) and not (destcdt & housing_codes):
        return {
            "habitat_eligible": False,
            "habitat_statut": "Exclue — logement/habitation interdit",
            "habitat_preuve": f"DESTNON={','.join(sorted(destnon & housing_codes))}",
            "habitat_confiance": 100,
            "typezone_normalise": typezone,
            "economic_vocation": False,
        }

    if destoui & housing_codes:
        return {
            "habitat_eligible": True,
            "habitat_statut": "Habitat autorisé",
            "habitat_preuve": f"DESTOUI={','.join(sorted(destoui & housing_codes))}",
            "habitat_confiance": 100,
            "typezone_normalise": typezone,
            "economic_vocation": False,
        }

    if destcdt & housing_codes:
        if include_conditionnel:
            return {
                "habitat_eligible": True,
                "habitat_statut": "Habitat autorisé sous conditions",
                "habitat_preuve": f"DESTCDT={','.join(sorted(destcdt & housing_codes))}",
                "habitat_confiance": 90,
                "typezone_normalise": typezone,
                "economic_vocation": False,
            }
        return {
            "habitat_eligible": False,
            "habitat_statut": "Exclue — habitat seulement sous conditions",
            "habitat_preuve": f"DESTCDT={','.join(sorted(destcdt & housing_codes))}",
            "habitat_confiance": 90,
            "typezone_normalise": typezone,
            "economic_vocation": False,
        }

    # Si les champs DEST* sont renseignés mais ne contiennent ni Habitation ni Logement,
    # on considère qu'il n'y a pas de preuve de destination logement.
    if destoui or destcdt or destnon:
        return {
            "habitat_eligible": False,
            "habitat_statut": "Exclue — destination logement non identifiée",
            "habitat_preuve": (
                f"DESTOUI={','.join(sorted(destoui)) or '—'} ; "
                f"DESTCDT={','.join(sorted(destcdt)) or '—'} ; "
                f"DESTNON={','.join(sorted(destnon)) or '—'}"
            ),
            "habitat_confiance": 95,
            "typezone_normalise": typezone,
            "economic_vocation": False,
        }

    # 2. Ancien standard CNIG : vocation dominante.
    # 01 = habitat ; 03 = mixte habitat/activité.
    dd = re.sub(r"\D", "", destdomi)
    if dd in {"1", "01"}:
        return {
            "habitat_eligible": True,
            "habitat_statut": "Habitat — vocation dominante",
            "habitat_preuve": "DESTDOMI=01 (habitat)",
            "habitat_confiance": 95,
            "typezone_normalise": typezone,
            "economic_vocation": False,
        }
    if dd in {"3", "03"}:
        return {
            "habitat_eligible": True,
            "habitat_statut": "Habitat/mixte — vocation dominante",
            "habitat_preuve": "DESTDOMI=03 (mixte habitat/activité)",
            "habitat_confiance": 90,
            "typezone_normalise": typezone,
            "economic_vocation": False,
        }
    if dd in {"2", "02", "4", "04", "5", "05", "7", "07", "8", "08", "9", "09", "10"}:
        return {
            "habitat_eligible": False,
            "habitat_statut": "Exclue — vocation dominante non habitat",
            "habitat_preuve": f"DESTDOMI={destdomi}",
            "habitat_confiance": 95,
            "typezone_normalise": typezone,
            "economic_vocation": False,
        }

    # 3. Analyse sémantique du libellé et du nom long.
    positive = any(k in text for k in HABITAT_KEYWORDS)
    negative = any(k in text for k in NON_RESIDENTIAL_KEYWORDS)

    if positive and not negative:
        return {
            "habitat_eligible": True,
            "habitat_statut": "Habitat probable — libellé PLU explicite",
            "habitat_preuve": f"{libelle} — {libelong}".strip(" —"),
            "habitat_confiance": 80,
            "typezone_normalise": typezone,
            "economic_vocation": False,
        }

    if negative and not positive:
        return {
            "habitat_eligible": False,
            "habitat_statut": "Exclue — libellé PLU non résidentiel",
            "habitat_preuve": f"{libelle} — {libelong}".strip(" —"),
            "habitat_confiance": 80,
            "typezone_normalise": typezone,
            "economic_vocation": False,
        }

    # 4. Haute précision : en mode Habitat strict, une zone sans preuve explicite
    # n'est pas retenue. Cela réduit les faux positifs.
    return {
        "habitat_eligible": False,
        "habitat_statut": "Exclue — destination habitat non prouvée",
        "habitat_preuve": f"{libelle} — {libelong}".strip(" —") or "Métadonnées PLU insuffisantes",
        "habitat_confiance": 60,
        "typezone_normalise": typezone,
    }


def classify_zone(props):
    """Compatibilité avec quelques affichages historiques de l'application."""
    a = analyse_habitat_zone(props, include_au=True, include_conditionnel=True)
    return a["habitat_statut"], a["habitat_eligible"]


def parcel_properties(feature):
    p = feature.get("properties", {}) or {}
    section = str(p.get("section") or "").strip()
    numero = str(p.get("numero") or "").strip()
    prefixe = str(p.get("prefixe") or "").strip()
    parcel_id = str(feature.get("id") or p.get("id") or "").strip()
    contenance = p.get("contenance")
    try:
        contenance = float(contenance)
    except Exception:
        contenance = None
    return {
        "section": section,
        "numero": numero,
        "prefixe": prefixe,
        "id_parcelle": parcel_id,
        "contenance": contenance,
    }


INDUSTRIAL_STRICT_KEYWORDS = [
    "industrie",
    "industriel",
    "industrielle",
    "industriels",
    "industrielles",
    "zone industrielle",
    "secteur industriel",
    "activite industrielle",
    "activites industrielles",
    "logistique",
    "plateforme logistique",
    "entrepot",
    "entrepots",
    "artisanale",
    "zone artisanale",
    "parc d activites",
    "zone d activites",
    "zone d activite",
    "vocation economique",
    "activites economiques",
]


def classify_zone_for_search(props, include_au=True):
    """
    Filtre métier V3.8 :
    - exclut A, N, AUs et assimilées ;
    - exclut les zones d'activités économiques/industrielles ;
    - conserve U et, si demandé, AUc/AU ouverts à l'urbanisation.

    Le filtre ne cherche PAS à imposer une destination habitat : l'utilisateur a
    demandé de supprimer ce filtre. Il élimine uniquement les zones manifestement
    non constructibles ou économiques/industrielles.
    """
    typezone = _normalise_typezone(props.get("typezone"))
    tu = normalize_text(typezone).replace(" ", "")
    libelle = str(props.get("libelle") or "")
    libelong = str(props.get("libelong") or "")
    text = normalize_text(f"{libelle} {libelong}")

    economic, economic_reason = detect_economic_vocation(props)

    # Anciennes variantes Ah/Nh/Nd : elles restent des zones A/N.
    if tu.startswith("a") and tu not in {"au", "auc", "aus"}:
        return False, "Zone agricole / non constructible", typezone
    if tu.startswith("n"):
        return False, "Zone naturelle / non constructible", typezone
    if typezone == "AUs":
        return False, "Zone AU bloquée (AUs)", typezone

    if typezone in {"AUc", "AU"} and not include_au:
        return False, "Zone AU désactivée dans les critères", typezone

    if typezone not in {"U", "AUc", "AU"}:
        return False, f"Type de zone non retenu : {typezone or 'inconnu'}", typezone

    if economic:
        return False, economic_reason, typezone

    if any(k in text for k in INDUSTRIAL_STRICT_KEYWORDS):
        return False, f"Zone industrielle/économique : {libelle or libelong}", typezone

    # Certains PLU anciens ne renseignent ni FORMDOMI ni un LIBELONG explicite,
    # mais utilisent des codes locaux très classiques pour les zones d'activités.
    # On applique cette heuristique seulement aux codes les plus caractéristiques.
    zone_code = re.sub(r"[^A-Z0-9]", "", str(libelle or "").upper())
    if re.match(r"^(?:U(?:X|Y|I)[A-Z0-9]*|(?:1|2)?AU(?:X|Y|I)[A-Z0-9]*)$", zone_code):
        return False, f"Code de zone d'activités/industrie : {libelle}", typezone

    return True, "Zone constructible retenue", typezone


def _rule_text_norm(value):
    txt = str(value or "").replace("\xa0", " ")
    txt = unicodedata.normalize("NFKD", txt)
    txt = "".join(c for c in txt if not unicodedata.combining(c))
    txt = txt.lower().replace("’", "'").replace("–", "-").replace("—", "-")
    return re.sub(r"\s+", " ", txt).strip()


def _url_page_hint(url):
    try:
        fragment = urlparse(str(url or "")).fragment or ""
        m = re.search(r"(?:^|&)page=(\d+)", fragment, flags=re.I)
        return int(m.group(1)) if m else None
    except Exception:
        return None


@st.cache_data(ttl=24 * 3600, show_spinner=False)
def fetch_rule_pdf_pages(url):
    """
    Télécharge uniquement le règlement lié à une zone/prescription.
    Pas d'archive PLUi complète : cette approche reste compatible avec Streamlit.
    """
    if not url:
        return [], None

    r = requests.get(url, headers=HEADERS, timeout=75)
    r.raise_for_status()

    if r.content[:4] != b"%PDF" and "application/pdf" not in (
        r.headers.get("content-type") or ""
    ).lower():
        raise RuntimeError("Le lien URLFIC ne renvoie pas directement un PDF.")

    reader = PdfReader(io.BytesIO(r.content))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        pages.append(
            {
                "page": i + 1,
                "text": txt,
                "norm": _rule_text_norm(txt),
            }
        )

    return pages, _url_page_hint(url)


def _zone_rule_pages(pages, zone_label, page_hint=None):
    """
    Sélectionne uniquement les pages susceptibles de concerner la zone.
    URLFIC#page=N est prioritaire. Sinon on repère les titres de zone.
    """
    if not pages:
        return []

    indexes = set()

    if page_hint:
        idx = max(0, int(page_hint) - 1)
        for j in range(max(0, idx - 2), min(len(pages), idx + 14)):
            indexes.add(j)

    z = _rule_text_norm(zone_label)
    if z:
        scored = []
        for i, p in enumerate(pages):
            txt = p["norm"]
            score = 0
            if re.search(rf"\bzone\s+{re.escape(z)}\b", txt):
                score += 8
            elif re.search(rf"\bsecteur\s+{re.escape(z)}\b", txt):
                score += 7
            elif len(z) >= 2 and re.search(rf"\b{re.escape(z)}\b", txt):
                score += 2

            if "emprise au sol" in txt:
                score += 2
            if "hauteur" in txt or re.search(r"\br\s*\+\s*\d+", txt):
                score += 2

            if score >= 5:
                scored.append((score, i))

        if scored:
            best = max(s for s, _ in scored)
            for score, i in scored:
                if score >= best - 1:
                    for j in range(max(0, i - 2), min(len(pages), i + 12)):
                        indexes.add(j)

    # Si le document est court, il s'agit souvent du règlement de zone lui-même.
    if not indexes and len(pages) <= 45:
        indexes.update(range(len(pages)))

    return [pages[i] for i in sorted(indexes)]


def _rule_context_score(text):
    t = _rule_text_norm(text)
    score = 0

    for x in [
        "maximum",
        "maximale",
        "ne peut exceder",
        "ne doit pas exceder",
        "au maximum",
        "est limitee a",
        "est limite a",
        "fixee a",
        "fixe a",
    ]:
        if x in t:
            score += 2

    # Éviter de prendre les règles d'annexes, clôtures, piscines, etc.
    for x in [
        "annexe",
        "abri",
        "piscine",
        "cloture",
        "local technique",
        "equipement public",
    ]:
        if x in t:
            score -= 2

    return score


def extract_emprise_from_rule_text(text):
    t = _rule_text_norm(text)
    candidates = []

    for m in re.finditer(
        r"(?:emprise\s+au\s+sol|coefficient\s+d[' ]?emprise|\bces\b)",
        t,
    ):
        w = t[max(0, m.start() - 260): min(len(t), m.end() + 720)]

        if "non reglement" in w or "sans objet" in w:
            continue

        for p in re.finditer(r"(?<!\d)(\d{1,3}(?:[.,]\d+)?)\s*%", w):
            val = float(p.group(1).replace(",", "."))
            if 0 < val <= 100:
                around = w[max(0, p.start() - 180): p.end() + 180]
                score = _rule_context_score(around)
                if "emprise" in around or "ces" in around:
                    score += 6
                candidates.append((score, val, around))

        for p in re.finditer(
            r"(?:ces|coefficient\s+d[' ]?emprise).{0,120}?"
            r"(?:=|:|fixe(?:e)?\s+a)?\s*(0[.,]\d+)",
            w,
        ):
            val = float(p.group(1).replace(",", ".")) * 100
            if 0 < val <= 100:
                around = w[max(0, p.start() - 170): p.end() + 190]
                candidates.append((_rule_context_score(around) + 7, val, around))

    if not candidates:
        return None, 0, ""

    candidates.sort(key=lambda x: (-x[0], x[1]))
    score, value, excerpt = candidates[0]
    confidence = 93 if score >= 10 else 84 if score >= 7 else 68
    return round(value, 2), confidence, excerpt[:750]


def extract_height_from_rule_text(text, floor_height_m):
    t = _rule_text_norm(text)
    direct = []
    heights = []

    for m in re.finditer(r"(?:hauteur|niveaux?|r\s*\+\s*\d+)", t):
        w = t[max(0, m.start() - 280): min(len(t), m.end() + 850)]

        # Priorité à R+N ou à un nombre de niveaux explicite.
        for p in re.finditer(r"\br\s*\+\s*(\d{1,2})\b", w):
            n = int(p.group(1)) + 1
            if 1 <= n <= 20:
                around = w[max(0, p.start() - 180): p.end() + 220]
                direct.append(
                    (_rule_context_score(around) + 9, n, None, "R+N explicite", around)
                )

        for pat in [
            r"(?<!\d)(\d{1,2})\s+niveaux?\s*(?:maximum|maxi)?",
            r"(?:maximum|maximale?)\s+(?:de\s+)?(\d{1,2})\s+niveaux?",
        ]:
            for p in re.finditer(pat, w):
                n = int(p.group(1))
                if 1 <= n <= 20:
                    around = w[max(0, p.start() - 180): p.end() + 220]
                    direct.append(
                        (_rule_context_score(around) + 7, n, None, "niveaux explicites", around)
                    )

        # Hauteur en mètres.
        for p in re.finditer(
            r"(?:hauteur.{0,180}?)(\d{1,2}(?:[.,]\d+)?)\s*m(?:etre)?s?\b",
            w,
        ):
            h = float(p.group(1).replace(",", "."))
            if 2.5 <= h <= 60:
                around = w[max(0, p.start() - 180): p.end() + 220]
                score = _rule_context_score(around) + 3
                if any(x in around for x in ["egout", "acrot", "facade"]):
                    score += 4
                if "faitage" in around and not any(
                    x in around for x in ["egout", "acrot"]
                ):
                    score -= 2
                heights.append((score, h, around))

    if direct:
        direct.sort(key=lambda x: (-x[0], -x[1]))
        score, levels, _, method, excerpt = direct[0]
        conf = 96 if score >= 12 else 87 if score >= 8 else 74
        return levels, None, conf, method, excerpt[:750]

    if heights:
        heights.sort(key=lambda x: (-x[0], x[1]))
        score, h, excerpt = heights[0]
        levels = max(
            1,
            int(math.floor(float(h) / max(float(floor_height_m), 0.1))),
        )
        conf = 78 if score >= 10 else 67 if score >= 7 else 54
        return levels, round(h, 2), conf, "hauteur convertie en niveaux", excerpt[:750]

    return None, None, 0, "", ""


@st.cache_data(ttl=24 * 3600, show_spinner=False)
def extract_zone_base_rule(reglement_url, zone_label, floor_height_m):
    """
    Valeur de base attachée à la zone graphique.
    Les prescriptions 38/39 locales restent prioritaires ensuite.
    """
    empty = {
        "emprise_pct": None,
        "levels": None,
        "height_m": None,
        "emprise_conf": 0,
        "height_conf": 0,
        "emprise_source": "",
        "height_source": "",
        "emprise_excerpt": "",
        "height_excerpt": "",
    }

    if not reglement_url:
        return empty

    try:
        pages, page_hint = fetch_rule_pdf_pages(reglement_url)
        subset = _zone_rule_pages(
            pages,
            zone_label=zone_label,
            page_hint=page_hint,
        )
        if not subset:
            return empty

        text = "\n".join(p["text"] for p in subset)

        emprise, ec, ee = extract_emprise_from_rule_text(text)
        levels, height, hc, hm, he = extract_height_from_rule_text(
            text,
            floor_height_m=floor_height_m,
        )

        return {
            "emprise_pct": emprise,
            "levels": levels,
            "height_m": height,
            "emprise_conf": ec,
            "height_conf": hc,
            "emprise_source": "Règlement lié à la zone graphique" if emprise is not None else "",
            "height_source": (
                f"Règlement lié à la zone graphique — {hm}"
                if levels is not None
                else ""
            ),
            "emprise_excerpt": ee,
            "height_excerpt": he,
        }
    except Exception:
        return empty


def build_commune_zone_rule_catalog(zones_geojson, include_au, floor_height_m):
    """
    Pré-interprétation SPÉCIFIQUE à la commune sélectionnée.
    Une règle n'est lue qu'une fois par zone/règlement.
    """
    catalog = {}
    for zf in zones_geojson.get("features", []) or []:
        props = zf.get("properties", {}) or {}
        allowed, reason, typezone = classify_zone_for_search(
            props,
            include_au=include_au,
        )
        if not allowed:
            continue

        label = str(props.get("libelle") or "")
        url = str(props.get("urlfic") or "")
        key = (label, url)
        if key in catalog:
            continue

        catalog[key] = extract_zone_base_rule(
            url,
            label,
            float(floor_height_m),
        )

    return catalog


def _overlay_piece_value(pieces, rule_geom, field, value, source, confidence):
    """
    Découpe une sous-surface de parcelle. La prescription graphique ne s'applique
    qu'à son périmètre réel, jamais à toute la parcelle par approximation.
    """
    out = []
    for piece in pieces:
        geom = piece["geom"]
        try:
            inter = geom.intersection(rule_geom)
            diff = geom.difference(rule_geom)
        except Exception:
            out.append(piece)
            continue

        if not inter.is_empty and geodesic_area_m2(inter) > 0.05:
            p = dict(piece)
            p["geom"] = inter
            p[field] = value
            p[field + "_source"] = source
            p[field + "_conf"] = confidence
            out.append(p)

        if not diff.is_empty and geodesic_area_m2(diff) > 0.05:
            p = dict(piece)
            p["geom"] = diff
            out.append(p)

    return out


def build_nonbuildable_graphic_index(graphic_prescriptions):
    """
    CNIG 02-01 = secteur avec interdiction de constructibilité.
    Ces surfaces sont retranchées de la partie constructible de la parcelle.
    """
    rules = []
    for feat in (
        (graphic_prescriptions or {})
        .get("prescription-surf", {})
        .get("features", [])
        or []
    ):
        props = feat.get("properties", {}) or {}
        t = _clean_psc_code(props.get("typepsc"))
        st = _clean_psc_code(props.get("stypepsc"))
        if t != "02" or st != "01":
            continue

        geom = valid_shape(feat)
        if geom is None:
            continue

        rules.append(
            {
                "geom": geom,
                "libelle": str(props.get("libelle") or ""),
                "txt": str(props.get("txt") or ""),
                "urlfic": str(props.get("urlfic") or ""),
            }
        )

    geoms = [r["geom"] for r in rules]
    return {
        "rules": rules,
        "tree": STRtree(geoms) if geoms else None,
    }


def subtract_nonbuildable_graphics(geom, nonbuildable_index):
    """
    Soustrait les secteurs graphiques explicitement inconstructibles (02-01).
    """
    if geom is None or geom.is_empty:
        return geom, 0.0, []

    tree = (nonbuildable_index or {}).get("tree")
    rules = (nonbuildable_index or {}).get("rules", [])
    if tree is None:
        return geom, 0.0, []

    cuts = []
    labels = []

    for ridx in tree.query(geom, predicate="intersects"):
        rule = rules[int(ridx)]
        try:
            inter = geom.intersection(rule["geom"])
        except Exception:
            continue
        if inter.is_empty:
            continue
        if geodesic_area_m2(inter) <= 0.05:
            continue
        cuts.append(inter)
        labels.append(rule.get("libelle") or rule.get("txt") or "CNIG 02-01")

    if not cuts:
        return geom, 0.0, []

    cut_union = unary_union(cuts)
    try:
        remaining = geom.difference(cut_union)
    except Exception:
        return geom, 0.0, []

    removed = geodesic_area_m2(geom.intersection(cut_union))
    return remaining, removed, list(dict.fromkeys(labels))


def enrich_with_precise_commune_rules(
    df,
    feature_map,
    rule_indexes,
    zone_rule_catalog,
    ratio_sdp_pct,
    ratio_shab_pct,
    shab_par_logement,
    floor_height_m,
):
    """
    Calcul V3.8 :
    1. la parcelle est déjà découpée sur les ZONE_URBA officielles ;
    2. chaque morceau reçoit la règle de base de sa zone ;
    3. les prescriptions graphiques 38-02 / 39-02 remplacent localement la règle ;
    4. surface brute = somme des sous-surfaces × emprise × niveaux.

    Aucune zone A/N/AUs/économique/industrielle ne participe au calcul.
    """
    df = df.copy()

    defaults = {
        "emprise_plu_pct": None,
        "hauteur_plu_m": None,
        "niveaux_plu": None,
        "surface_brute_plu_m2": None,
        "sdp_plu_m2": None,
        "shab_plu_m2": None,
        "logements_plu": None,
        "gabarit_plu_statut": "À vérifier",
        "gabarit_plu_source": "",
        "gabarit_plu_confiance": 0,
        "gabarit_couverture_pct": 0.0,
        "prescription_emprise": "",
        "prescription_hauteur": "",
        "regle_emprise_source": "",
        "regle_hauteur_source": "",
        "regle_emprise_extrait": "",
        "regle_hauteur_extrait": "",
    }
    for c, default in defaults.items():
        if c not in df.columns:
            df[c] = default

    if df.empty:
        return df

    surf_tree = rule_indexes.get("surf_tree")
    surf_rules = rule_indexes.get("surf_rules", [])

    for idx, row in df.iterrows():
        feat = feature_map.get(row["reference"])
        if not feat:
            continue

        zone_pieces = feat.get("_zone_pieces", []) or []
        if not zone_pieces:
            continue

        all_pieces = []
        applied_e = []
        applied_h = []

        for zp in zone_pieces:
            zg = shape(zp["geometry"])
            key = (str(zp.get("zone_plu") or ""), str(zp.get("reglement_url") or ""))
            base = zone_rule_catalog.get(key, {})

            base_e = base.get("emprise_pct")
            base_l = base.get("levels")

            pieces = [
                {
                    "geom": zg,
                    "emprise": base_e,
                    "levels": base_l,
                    "emprise_source": base.get("emprise_source", ""),
                    "levels_source": base.get("height_source", ""),
                    "emprise_conf": base.get("emprise_conf", 0),
                    "levels_conf": base.get("height_conf", 0),
                    "height_m": base.get("height_m"),
                    "emprise_excerpt": base.get("emprise_excerpt", ""),
                    "height_excerpt": base.get("height_excerpt", ""),
                }
            ]

            # Prescriptions surfaciques 38/39 : source graphique la plus précise.
            if surf_tree is not None:
                hits = []
                for ridx in surf_tree.query(zg, predicate="intersects"):
                    rule = surf_rules[int(ridx)]
                    try:
                        inter = zg.intersection(rule["geom"])
                        area = geodesic_area_m2(inter)
                    except Exception:
                        area = 0.0
                    if area > 0.05:
                        hits.append((area, rule))

                # Les plus grandes prescriptions sont appliquées d'abord ;
                # les plus petites/locales appliquées ensuite les remplacent localement.
                hits.sort(key=lambda x: -x[0])

                for _, rule in hits:
                    if (
                        rule.get("typepsc") == "38"
                        and rule.get("emprise_pct") is not None
                    ):
                        pieces = _overlay_piece_value(
                            pieces,
                            rule["geom"],
                            "emprise",
                            float(rule["emprise_pct"]),
                            f"Prescription graphique {rule.get('typepsc')}-{rule.get('stypepsc')}",
                            int(rule.get("confidence", 0)),
                        )
                        applied_e.append(rule)

                    if rule.get("typepsc") == "39":
                        levels = rule.get("levels")
                        height = rule.get("height_m")
                        if levels is None and height is not None:
                            levels = max(
                                1,
                                int(
                                    math.floor(
                                        float(height)
                                        / max(float(floor_height_m), 0.1)
                                    )
                                ),
                            )
                        if levels is not None:
                            pieces = _overlay_piece_value(
                                pieces,
                                rule["geom"],
                                "levels",
                                int(levels),
                                f"Prescription graphique {rule.get('typepsc')}-{rule.get('stypepsc')}",
                                int(rule.get("confidence", 0)),
                            )
                            applied_h.append(rule)

            all_pieces.extend(pieces)

        eligible_area = max(
            float(row.get("surface_plu_eligible_m2") or 0),
            0.01,
        )
        covered_area = 0.0
        gross = 0.0
        weighted_e = 0.0
        weighted_levels = 0.0
        weighted_height = 0.0
        height_area = 0.0
        confs = []
        emprise_sources = []
        height_sources = []
        emprise_excerpts = []
        height_excerpts = []

        for piece in all_pieces:
            area = geodesic_area_m2(piece["geom"])
            e = piece.get("emprise")
            levels = piece.get("levels")

            if e is not None:
                weighted_e += area * float(e)
            if levels is not None:
                weighted_levels += area * float(levels)
            if piece.get("height_m") is not None:
                weighted_height += area * float(piece["height_m"])
                height_area += area

            if piece.get("emprise_source"):
                emprise_sources.append(piece["emprise_source"])
            if piece.get("levels_source"):
                height_sources.append(piece["levels_source"])
            if piece.get("emprise_excerpt"):
                emprise_excerpts.append(piece["emprise_excerpt"])
            if piece.get("height_excerpt"):
                height_excerpts.append(piece["height_excerpt"])

            if e is not None and levels is not None:
                covered_area += area
                gross += area * (float(e) / 100.0) * float(levels)
                if piece.get("emprise_conf"):
                    confs.append(int(piece["emprise_conf"]))
                if piece.get("levels_conf"):
                    confs.append(int(piece["levels_conf"]))

        coverage = max(0.0, min(100.0, 100.0 * covered_area / eligible_area))
        df.at[idx, "gabarit_couverture_pct"] = round(coverage, 1)

        if weighted_e > 0:
            df.at[idx, "emprise_plu_pct"] = round(weighted_e / eligible_area, 2)
        if weighted_levels > 0:
            df.at[idx, "niveaux_plu"] = round(weighted_levels / eligible_area, 2)
        if height_area > 0:
            df.at[idx, "hauteur_plu_m"] = round(weighted_height / height_area, 2)

        # La capacité n'est validée automatiquement que si presque toute la partie
        # constructible de la parcelle dispose d'une emprise ET d'une hauteur.
        if coverage >= 95.0 and gross > 0:
            sdp = gross * float(ratio_sdp_pct) / 100.0
            shab = sdp * float(ratio_shab_pct) / 100.0
            logements = math.floor(
                shab / max(float(shab_par_logement), 1.0)
            )

            df.at[idx, "surface_brute_plu_m2"] = round(gross)
            df.at[idx, "sdp_plu_m2"] = round(sdp)
            df.at[idx, "shab_plu_m2"] = round(shab)
            df.at[idx, "logements_plu"] = int(max(0, logements))
            df.at[idx, "gabarit_plu_statut"] = "Calculé sur cadastre + PLU/PLUi en vigueur"
            df.at[idx, "gabarit_plu_source"] = (
                "Zonage graphique + prescriptions 38/39 + règlement de zone"
            )
            df.at[idx, "gabarit_plu_confiance"] = (
                min(confs) if confs else 60
            )
        else:
            missing = []
            if df.at[idx, "emprise_plu_pct"] in [None, ""]:
                missing.append("emprise non chiffrée")
            if df.at[idx, "niveaux_plu"] in [None, ""]:
                missing.append("hauteur/niveaux non chiffrés")
            if coverage < 95:
                missing.append(f"couverture règles {coverage:.0f}%")
            df.at[idx, "gabarit_plu_statut"] = "À vérifier — " + ", ".join(missing)
            df.at[idx, "gabarit_plu_source"] = "PLU/PLUi spécifique commune — données incomplètes"

        if applied_e:
            df.at[idx, "prescription_emprise"] = " / ".join(
                dict.fromkeys(
                    f"{r.get('typepsc')}-{r.get('stypepsc')} "
                    f"{r.get('libelle') or r.get('txt') or ''}".strip()
                    for r in applied_e
                )
            )[:800]
        if applied_h:
            df.at[idx, "prescription_hauteur"] = " / ".join(
                dict.fromkeys(
                    f"{r.get('typepsc')}-{r.get('stypepsc')} "
                    f"{r.get('libelle') or r.get('txt') or ''}".strip()
                    for r in applied_h
                )
            )[:800]

        df.at[idx, "regle_emprise_source"] = " / ".join(
            dict.fromkeys(emprise_sources)
        )[:500]
        df.at[idx, "regle_hauteur_source"] = " / ".join(
            dict.fromkeys(height_sources)
        )[:500]
        df.at[idx, "regle_emprise_extrait"] = " | ".join(
            dict.fromkeys(emprise_excerpts)
        )[:1100]
        df.at[idx, "regle_hauteur_extrait"] = " | ".join(
            dict.fromkeys(height_excerpts)
        )[:1100]

    return df


def build_plu_commune_diagnostic(
    partition,
    documents,
    zones_geojson,
    graphic_prescriptions,
    include_au,
):
    zones = zones_geojson.get("features", []) or []
    retained = 0
    excluded_nonconstructible = 0
    excluded_economic = 0

    for zf in zones:
        props = zf.get("properties", {}) or {}
        allowed, reason, _ = classify_zone_for_search(props, include_au=include_au)
        if allowed:
            retained += 1
        else:
            r = normalize_text(reason)
            if "econom" in r or "industri" in r or "activite" in r:
                excluded_economic += 1
            else:
                excluded_nonconstructible += 1

    p38 = 0
    p39 = 0
    p02_01 = 0
    for layer, fc in (graphic_prescriptions or {}).items():
        for feat in fc.get("features", []) or []:
            props = feat.get("properties", {}) or {}
            t = _clean_psc_code(props.get("typepsc"))
            st = _clean_psc_code(props.get("stypepsc"))
            if t == "38" and st in {"00", "02"}:
                p38 += 1
            if t == "39" and st in {"00", "02", "50", "51"}:
                p39 += 1
            if t == "02" and st == "01":
                p02_01 += 1

    doc_props = {}
    if documents:
        doc_props = documents[0].get("properties", {}) or {}

    date_doc = (
        doc_props.get("datvalid")
        or doc_props.get("datappro")
        or doc_props.get("date_appro")
        or ""
    )
    idurba = doc_props.get("idurba") or doc_props.get("id") or ""
    typedoc = doc_props.get("typedoc") or doc_props.get("type_doc") or "PLU/PLUi"

    return {
        "partition": partition,
        "typedoc": typedoc,
        "idurba": idurba,
        "date_document": str(date_doc),
        "zones_total": len(zones),
        "zones_retenues": retained,
        "zones_non_constructibles_exclues": excluded_nonconstructible,
        "zones_economiques_exclues": excluded_economic,
        "prescriptions_emprise": p38,
        "prescriptions_hauteur": p39,
        "prescriptions_inconstructibles": p02_01,
    }


def analyse_commune(
    commune_name,
    insee,
    parcelles_geojson,
    batiments_geojson,
    zones_geojson,
    bdnb_parcel_index,
    nonbuildable_index,
    include_au,
    habitat_only,
    include_conditionnel,
    ratio_sdp_pct,
    ratio_shab_pct,
    shab_par_logement,
):
    """
    V3.8 : comparaison géométrique réelle du dernier cadastre avec le PLU/PLUi.

    Contrairement aux versions précédentes, on ne prend plus le point représentatif
    de la parcelle ni la première zone rencontrée. La parcelle est réellement
    intersectée avec toutes les zones du règlement graphique.
    """
    zone_rows = []
    zone_geoms = []

    for zf in zones_geojson.get("features", []) or []:
        zg = valid_shape(zf)
        if zg is None:
            continue

        props = zf.get("properties", {}) or {}
        allowed, exclusion_reason, typezone = classify_zone_for_search(
            props,
            include_au=include_au,
        )
        habitat = analyse_habitat_zone(
            props,
            include_au=include_au,
            include_conditionnel=include_conditionnel,
        )

        zone_rows.append(
            {
                "geom": zg,
                "allowed": bool(allowed),
                "exclusion_reason": exclusion_reason,
                "typezone": typezone,
                "libelle": props.get("libelle") or "",
                "libelong": props.get("libelong") or "",
                "destdomi": props.get("destdomi") or "",
                "formdomi": props.get("formdomi") or "",
                "economic_vocation": bool(habitat.get("economic_vocation", False)),
                "destoui": props.get("destoui") or "",
                "destcdt": props.get("destcdt") or "",
                "destnon": props.get("destnon") or "",
                "habitat_eligible": habitat.get("habitat_eligible", False),
                "habitat_statut": habitat.get("habitat_statut", ""),
                "habitat_preuve": habitat.get("habitat_preuve", ""),
                "habitat_confiance": habitat.get("habitat_confiance", 0),
                "url_reglement": props.get("urlfic") or "",
                "datvalid": props.get("datvalid") or props.get("datappro") or "",
            }
        )
        zone_geoms.append(zg)

    if not zone_geoms:
        return pd.DataFrame(), {}

    zone_tree = STRtree(zone_geoms)

    building_geoms = []
    building_centroids = []
    for bf in batiments_geojson.get("features", []) or []:
        bg = valid_shape(bf)
        if bg is None:
            continue
        building_geoms.append(bg)
        building_centroids.append(bg.representative_point())

    building_tree = STRtree(building_centroids) if building_centroids else None

    rows = []
    feature_map = {}

    for pf in parcelles_geojson.get("features", []) or []:
        pg = valid_shape(pf)
        if pg is None:
            continue

        parcel_area_geom = geodesic_area_m2(pg)
        if parcel_area_geom <= 0:
            continue

        zone_idxs = zone_tree.query(pg, predicate="intersects")
        if len(zone_idxs) == 0:
            continue

        eligible_pieces = []
        eligible_zone_pieces = []
        excluded_area = 0.0
        overlap_details = []

        for zi_raw in zone_idxs:
            zi = int(zi_raw)
            zr = zone_rows[zi]
            try:
                inter = pg.intersection(zr["geom"])
            except Exception:
                continue

            if inter.is_empty:
                continue

            inter_area = geodesic_area_m2(inter)
            if inter_area <= 0.05:
                continue

            overlap_details.append(
                (
                    inter_area,
                    zr["libelle"],
                    zr["typezone"],
                    zr["allowed"],
                    zr["exclusion_reason"],
                )
            )

            if zr["allowed"]:
                eligible_pieces.append(inter)
                eligible_zone_pieces.append(
                    {
                        "geometry": mapping(inter),
                        "zone_plu": zr["libelle"],
                        "zone_description": zr["libelong"],
                        "zone_type": zr["typezone"],
                        "reglement_url": zr["url_reglement"],
                        "date_zone": zr["datvalid"],
                        "area_m2": inter_area,
                        "economic_vocation": zr["economic_vocation"],
                    }
                )
            else:
                excluded_area += inter_area

        if not eligible_pieces:
            # Parcelle intégralement située en zone non constructible ou économique.
            continue

        eligible_geom = unary_union(eligible_pieces)

        # Retrait des secteurs graphiques explicitement inconstructibles CNIG 02-01.
        eligible_geom, removed_nonbuildable, nonbuildable_labels = (
            subtract_nonbuildable_graphics(
                eligible_geom,
                nonbuildable_index,
            )
        )
        eligible_area = geodesic_area_m2(eligible_geom)
        if eligible_area <= 0.05:
            continue

        # Re-découper les pièces de zones pour retirer également les secteurs 02-01.
        cleaned_zone_pieces = []
        for zp in eligible_zone_pieces:
            zg = shape(zp["geometry"])
            cleaned_geom, _, _ = subtract_nonbuildable_graphics(
                zg,
                nonbuildable_index,
            )
            if cleaned_geom is None or cleaned_geom.is_empty:
                continue
            a = geodesic_area_m2(cleaned_geom)
            if a <= 0.05:
                continue
            zpc = dict(zp)
            zpc["geometry"] = mapping(cleaned_geom)
            zpc["area_m2"] = a
            cleaned_zone_pieces.append(zpc)

        eligible_zone_pieces = cleaned_zone_pieces
        if not eligible_zone_pieces:
            continue

        # Zone dominante uniquement pour l'affichage ; le calcul garde toutes
        # les sous-zones constructibles séparément.
        eligible_zone_pieces.sort(key=lambda x: x["area_m2"], reverse=True)
        dominant = eligible_zone_pieces[0]

        pp = parcel_properties(pf)
        surface = pp["contenance"]
        if surface is None or surface <= 0:
            surface = parcel_area_geom

        built_count = 0
        built_footprint = 0.0
        if building_tree is not None:
            bidx = building_tree.query(pg, predicate="contains")
            built_count = len(bidx)
            if built_count:
                for bi in bidx:
                    built_footprint += geodesic_area_m2(
                        building_geoms[int(bi)]
                    )

        terrain_bati = built_count > 0

        raw_ref = pp["id_parcelle"] or f"{insee}-{pp['section']}-{pp['numero']}"
        ref_norm = _normalise_parcelle_id(raw_ref)
        bdnb_info = detect_collective_housing(
            bdnb_parcel_index.get(ref_norm, [])
        )

        # V3.8 : aucune capacité n'est inventée avant l'analyse emprise/hauteur.
        surface_brute = None
        sdp_estimee = None
        shab_estimee = None
        logements = None

        eligible_pct = min(
            100.0,
            100.0 * eligible_area / max(parcel_area_geom, 0.01),
        )
        excluded_pct = max(0.0, 100.0 - eligible_pct)

        # Score de prospection basé sur la qualité spatiale de la partie constructible.
        score = 55
        if eligible_pct >= 95:
            score += 15
        elif eligible_pct >= 75:
            score += 8
        if float(surface) >= 1500:
            score += 10
        if float(surface) >= 3000:
            score += 5
        if not terrain_bati:
            score += 5
        if dominant["zone_type"] in {"AUc", "AU"}:
            score -= 5
        score = max(0, min(100, score))

        rp = eligible_geom.representative_point()
        ref = raw_ref

        overlap_label = " / ".join(
            f"{z} {round(100*a/max(parcel_area_geom,0.01))}%"
            for a, z, _, allowed, _ in sorted(
                overlap_details,
                reverse=True,
                key=lambda x: x[0],
            )[:5]
        )

        rows.append(
            {
                "selection": False,
                "ville": commune_name,
                "code_insee": insee,
                "zone_ptz": zone_ptz_commune,
                "reference": ref,
                "section": pp["section"],
                "numero": pp["numero"],
                "surface_m2": round(surface),
                "surface_plu_eligible_m2": round(eligible_area),
                "surface_plu_exclue_m2": round(max(0.0, parcel_area_geom - eligible_area)),
                "part_constructible_plu_pct": round(eligible_pct, 1),
                "zones_intersectees": overlap_label,
                "surface_inconstructible_graphique_m2": round(removed_nonbuildable),
                "prescriptions_inconstructibles": " / ".join(nonbuildable_labels),
                "surface_brute_m2": None,
                "sdp_estimee_m2": None,
                "shab_estimee_m2": None,
                "ratio_sdp_pct": float(ratio_sdp_pct),
                "ratio_shab_pct": float(ratio_shab_pct),
                "shab_par_logement": float(shab_par_logement),
                "terrain_bati": terrain_bati,
                "nb_batiments": built_count,
                "emprise_batie_m2": round(built_footprint),
                "collectif_existant": bool(bdnb_info["collectif_existant"]),
                "usage_bdnb": bdnb_info["usage_bdnb"],
                "zone_type": dominant["zone_type"],
                "zone_plu": dominant["zone_plu"],
                "zone_description": dominant["zone_description"],
                "classe_zone": "Constructible — analyse géométrique PLU/PLUi",
                "habitat_eligible": True,
                "habitat_statut": "Filtre habitat désactivé",
                "habitat_preuve": "",
                "habitat_confiance": 0,
                "destoui": "",
                "destcdt": "",
                "destnon": "",
                "destdomi": "",
                "formdomi": "",
                "economic_vocation": False,
                "reglement_url": dominant["reglement_url"],
                "date_zone": dominant["date_zone"],
                "logements_estimes": None,
                "score": score,
                "latitude": rp.y,
                "longitude": rp.x,
                "adresse": bdnb_info["adresse_bdnb"],
            }
        )

        feature_map[ref] = {
            "type": "Feature",
            "properties": {
                "reference": ref,
                "section": pp["section"],
                "numero": pp["numero"],
                "surface_m2": round(surface),
                "surface_constructible_plu": round(eligible_area),
                "part_constructible_plu": round(eligible_pct, 1),
                "surface_inconstructible_graphique": round(removed_nonbuildable),
                "zone": dominant["zone_plu"],
                "typezone": dominant["zone_type"],
                "collectif": "Oui" if bdnb_info["collectif_existant"] else "Non",
                "zones_intersectees": overlap_label,
            },
            "geometry": mapping(pg),
            # Données internes d'analyse ; pydeck ignore ces clés.
            "_analysis_geometry": mapping(eligible_geom),
            "_zone_pieces": eligible_zone_pieces,
        }

    expected_columns = [
        "selection",
        "ville",
        "code_insee",
        "zone_ptz",
        "reference",
        "section",
        "numero",
        "surface_m2",
        "surface_plu_eligible_m2",
        "surface_plu_exclue_m2",
        "part_constructible_plu_pct",
        "zones_intersectees",
        "surface_inconstructible_graphique_m2",
        "prescriptions_inconstructibles",
        "surface_brute_m2",
        "sdp_estimee_m2",
        "shab_estimee_m2",
        "ratio_sdp_pct",
        "ratio_shab_pct",
        "shab_par_logement",
        "terrain_bati",
        "nb_batiments",
        "emprise_batie_m2",
        "collectif_existant",
        "usage_bdnb",
        "zone_type",
        "zone_plu",
        "zone_description",
        "classe_zone",
        "habitat_eligible",
        "habitat_statut",
        "habitat_preuve",
        "habitat_confiance",
        "destoui",
        "destcdt",
        "destnon",
        "destdomi",
        "formdomi",
        "economic_vocation",
        "reglement_url",
        "date_zone",
        "logements_estimes",
        "score",
        "latitude",
        "longitude",
        "adresse",
        "proprietaire_personne_morale",
        "forme_juridique_proprietaire",
        "siren_proprietaire",
        "proprietaire_type",
        "proprietaire_commune",
    ]

    df = pd.DataFrame(rows)
    if df.empty:
        df = pd.DataFrame(columns=expected_columns)
    else:
        for col in expected_columns:
            if col not in df.columns:
                df[col] = None
        df = df[expected_columns]

    return df, feature_map


# --------------------------
# Courriers + dossier PDF
# --------------------------
def replace_text_in_runs(paragraph, replacements):
    """
    Remplacement robuste tout en conservant la mise en forme du premier run.
    Les placeholders du modèle peuvent être découpés sur plusieurs runs Word.
    """
    full = "".join(run.text for run in paragraph.runs)
    if not full:
        return

    new = full
    for old, repl in replacements.items():
        if old in new:
            new = new.replace(old, str(repl))

    if new != full:
        if paragraph.runs:
            paragraph.runs[0].text = new
            for r in paragraph.runs[1:]:
                r.text = ""
        else:
            paragraph.text = new


def _apply_docx_replacements(doc, replacements):
    for p in doc.paragraphs:
        replace_text_in_runs(p, replacements)

    for table in doc.tables:
        for row_ in table.rows:
            for cell in row_.cells:
                for p in cell.paragraphs:
                    replace_text_in_runs(p, replacements)

    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_text_in_runs(p, replacements)
        for p in section.footer.paragraphs:
            replace_text_in_runs(p, replacements)


def generate_letter(row, signataire, fonction, email, ville_signature):
    """
    Génère la lettre à partir du nouveau modèle SAGEC fourni par l'utilisateur.
    Le texte du modèle est conservé ; seuls les champs variables sont remplacés.
    """
    if not LETTER_TEMPLATE.exists():
        raise FileNotFoundError(
            "Le fichier lettre_sagec_modele.docx doit être présent à côté de app.py."
        )

    doc = Document(LETTER_TEMPLATE)

    adresse = str(row.get("adresse") or "").strip()
    if not adresse:
        adresse = (
            f"Parcelle section {row.get('section', '')} "
            f"n° {row.get('numero', '')}, {row.get('ville', '')}"
        )

    ville = str(row.get("ville") or "").strip()
    section = str(row.get("section") or "").strip()
    numero = str(row.get("numero") or "").strip()

    replacements = {
        "(Adresse)": adresse,
        "(Ville)": ville,
        "Anglet, le ../../2026.": (
            f"{ville_signature}, le {date.today().strftime('%d/%m/%Y')}."
        ),
        "Objet : Votre propriété située à …………….": (
            f"Objet : Votre propriété située à {adresse}"
        ),
        "votre propriété située à ………, cadastrée section ………………………,": (
            f"votre propriété située à {adresse}, "
            f"cadastrée section {section} n° {numero},"
        ),
        "Nicolas PEDROT": signataire,
        "Directeur": fonction,
        "Nicolas.pedrot@sagec.fr": email,
    }

    _apply_docx_replacements(doc, replacements)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _pdf_escape(value):
    txt = str(value if value not in [None, ""] else "—")
    return (
        txt.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _pdf_number(value, digits=0, suffix=""):
    try:
        if value is None or pd.isna(value):
            return "—"
        num = float(value)
        if digits == 0:
            text = f"{num:,.0f}".replace(",", " ")
        else:
            text = f"{num:,.{digits}f}".replace(",", " ").replace(".", ",")
        return f"{text}{suffix}"
    except Exception:
        return "—"


def _logo_bytes_from_template():
    """
    Réutilise directement le logo contenu dans le modèle Word SAGEC :
    aucun fichier logo supplémentaire n'est nécessaire sur GitHub.
    """
    if not LETTER_TEMPLATE.exists():
        return None
    try:
        with zipfile.ZipFile(LETTER_TEMPLATE, "r") as z:
            media = [
                n for n in z.namelist()
                if n.lower().startswith("word/media/")
                and n.lower().endswith((".png", ".jpg", ".jpeg"))
            ]
            if not media:
                return None
            media.sort(key=lambda n: len(z.read(n)), reverse=True)
            return z.read(media[0])
    except Exception:
        return None


class ParcelSketch(Flowable):
    """
    Petit plan vectoriel de la parcelle, dessiné directement depuis la géométrie cadastrale.
    """
    def __init__(self, feature, width=78 * mm, height=50 * mm):
        Flowable.__init__(self)
        self.feature = feature
        self.width = width
        self.height = height

    def draw(self):
        c = self.canv
        c.saveState()

        c.setStrokeColor(colors.HexColor("#D3DCE6"))
        c.setFillColor(colors.HexColor("#F7F9FB"))
        c.roundRect(0, 0, self.width, self.height, 4 * mm, fill=1, stroke=1)

        if not self.feature or not self.feature.get("geometry"):
            c.setFillColor(colors.HexColor("#7A8793"))
            c.setFont("Helvetica", 8)
            c.drawCentredString(
                self.width / 2,
                self.height / 2,
                "Géométrie cadastrale indisponible",
            )
            c.restoreState()
            return

        try:
            geom = shape(self.feature["geometry"])
            polygons = []
            if geom.geom_type == "Polygon":
                polygons = [geom]
            elif geom.geom_type == "MultiPolygon":
                polygons = list(geom.geoms)

            if not polygons:
                raise ValueError("géométrie non surfacique")

            minx, miny, maxx, maxy = geom.bounds
            dx = max(maxx - minx, 1e-9)
            dy = max(maxy - miny, 1e-9)

            pad = 5 * mm
            avail_w = self.width - 2 * pad
            avail_h = self.height - 2 * pad
            scale = min(avail_w / dx, avail_h / dy)

            offset_x = pad + (avail_w - dx * scale) / 2
            offset_y = pad + (avail_h - dy * scale) / 2

            c.setFillColor(colors.HexColor("#DCECF7"))
            c.setStrokeColor(colors.HexColor("#1766A5"))
            c.setLineWidth(1.4)

            for poly in polygons:
                coords = list(poly.exterior.coords)
                path = c.beginPath()
                for j, (x, y) in enumerate(coords):
                    px = offset_x + (x - minx) * scale
                    py = offset_y + (y - miny) * scale
                    if j == 0:
                        path.moveTo(px, py)
                    else:
                        path.lineTo(px, py)
                path.close()
                c.drawPath(path, fill=1, stroke=1)

                # Les trous cadastraux éventuels sont matérialisés en blanc.
                for ring in poly.interiors:
                    icoords = list(ring.coords)
                    ipath = c.beginPath()
                    for j, (x, y) in enumerate(icoords):
                        px = offset_x + (x - minx) * scale
                        py = offset_y + (y - miny) * scale
                        if j == 0:
                            ipath.moveTo(px, py)
                        else:
                            ipath.lineTo(px, py)
                    ipath.close()
                    c.setFillColor(colors.white)
                    c.drawPath(ipath, fill=1, stroke=1)
                    c.setFillColor(colors.HexColor("#DCECF7"))

            c.setFillColor(colors.HexColor("#5B6770"))
            c.setFont("Helvetica", 7)
            c.drawString(4 * mm, 2.5 * mm, "Emprise cadastrale - schéma non contractuel")

        except Exception:
            c.setFillColor(colors.HexColor("#7A8793"))
            c.setFont("Helvetica", 8)
            c.drawCentredString(
                self.width / 2,
                self.height / 2,
                "Aperçu parcellaire indisponible",
            )

        c.restoreState()


def generate_selected_parcels_pdf(selected_df, commune_name, insee, feature_map):
    """
    Génère un dossier PDF de prospection :
    - page de synthèse ;
    - tableau récapitulatif ;
    - une fiche détaillée par parcelle avec schéma cadastral.
    """
    buffer = io.BytesIO()
    page_w, page_h = A4

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=15 * mm,
        leftMargin=15 * mm,
        topMargin=31 * mm,
        bottomMargin=18 * mm,
        title=f"Sélection foncière - {commune_name}",
        author="SAGEC SUD ATLANTIQUE",
    )

    styles = getSampleStyleSheet()
    sagec_blue = colors.HexColor("#1766A5")
    dark = colors.HexColor("#263746")
    muted = colors.HexColor("#657786")
    light_blue = colors.HexColor("#EEF5FA")
    line = colors.HexColor("#D8E1E8")

    title_style = ParagraphStyle(
        "SagecTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=23,
        textColor=sagec_blue,
        spaceAfter=5 * mm,
    )
    subtitle_style = ParagraphStyle(
        "SagecSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        textColor=muted,
        spaceAfter=5 * mm,
    )
    h2 = ParagraphStyle(
        "SagecH2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=dark,
        spaceBefore=2 * mm,
        spaceAfter=3 * mm,
    )
    body = ParagraphStyle(
        "SagecBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=11,
        textColor=dark,
    )
    small = ParagraphStyle(
        "SagecSmall",
        parent=body,
        fontSize=7.3,
        leading=9,
        textColor=muted,
    )
    metric_label = ParagraphStyle(
        "MetricLabel",
        parent=body,
        fontSize=7,
        leading=8,
        textColor=muted,
        alignment=TA_CENTER,
    )
    metric_value = ParagraphStyle(
        "MetricValue",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=16,
        textColor=sagec_blue,
        alignment=TA_CENTER,
    )

    logo_bytes = _logo_bytes_from_template()
    logo_reader = None
    if logo_bytes:
        try:
            logo_reader = ImageReader(io.BytesIO(logo_bytes))
        except Exception:
            logo_reader = None

    def on_page(canvas, pdf_doc):
        canvas.saveState()

        # En-tête
        if logo_reader:
            try:
                iw, ih = logo_reader.getSize()
                max_w = 38 * mm
                max_h = 18 * mm
                ratio = min(max_w / iw, max_h / ih)
                canvas.drawImage(
                    logo_reader,
                    15 * mm,
                    page_h - 24 * mm,
                    width=iw * ratio,
                    height=ih * ratio,
                    mask="auto",
                )
            except Exception:
                pass

        canvas.setStrokeColor(line)
        canvas.setLineWidth(0.6)
        canvas.line(15 * mm, page_h - 27 * mm, page_w - 15 * mm, page_h - 27 * mm)

        # Pied de page
        canvas.setStrokeColor(line)
        canvas.line(15 * mm, 13 * mm, page_w - 15 * mm, 13 * mm)
        canvas.setFillColor(muted)
        canvas.setFont("Helvetica", 7)
        canvas.drawString(
            15 * mm,
            8 * mm,
            "SAGEC SUD ATLANTIQUE - Préselection foncière - Données indicatives",
        )
        canvas.drawRightString(
            page_w - 15 * mm,
            8 * mm,
            f"Page {pdf_doc.page}",
        )
        canvas.restoreState()

    rows = selected_df.copy()
    if "reference" in rows.columns:
        rows = rows.drop_duplicates(subset=["reference"]).reset_index(drop=True)

    story = []

    # --------------------------------------------------------------
    # Couverture / synthèse
    # --------------------------------------------------------------
    story.append(Paragraph("Sélection foncière", title_style))
    story.append(
        Paragraph(
            f"<b>{_pdf_escape(commune_name)}</b> - Code INSEE {_pdf_escape(insee)}"
            f"<br/>Dossier généré le {date.today().strftime('%d/%m/%Y')}",
            subtitle_style,
        )
    )

    total_surface = pd.to_numeric(rows.get("surface_m2"), errors="coerce").fillna(0).sum()
    total_log = pd.to_numeric(rows.get("logements_plu"), errors="coerce").fillna(0).sum()
    total_sdp = pd.to_numeric(rows.get("sdp_plu_m2"), errors="coerce").fillna(0).sum()

    ptz_values = []
    if "zone_ptz" in rows.columns:
        ptz_values = [
            str(x) for x in rows["zone_ptz"].dropna().unique().tolist() if str(x).strip()
        ]
    ptz_text = ", ".join(ptz_values) if ptz_values else "—"

    metrics = [
        (
            Paragraph(str(len(rows)), metric_value),
            Paragraph("parcelles sélectionnées", metric_label),
        ),
        (
            Paragraph(_pdf_number(total_surface, 0, " m²"), metric_value),
            Paragraph("surface cadastrale cumulée", metric_label),
        ),
        (
            Paragraph(_pdf_number(total_sdp, 0, " m²"), metric_value),
            Paragraph("SDP estimée cumulée", metric_label),
        ),
        (
            Paragraph(_pdf_number(total_log, 0), metric_value),
            Paragraph("logements estimés cumulés", metric_label),
        ),
    ]

    metric_cells = []
    for value, label in metrics:
        metric_cells.append(Table([[value], [label]], colWidths=[42 * mm]))

    metric_table = Table([metric_cells], colWidths=[43.5 * mm] * 4)
    metric_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), light_blue),
                ("BOX", (0, 0), (-1, -1), 0.6, line),
                ("INNERGRID", (0, 0), (-1, -1), 0.3, line),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2 * mm),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2 * mm),
                ("TOPPADDING", (0, 0), (-1, -1), 3 * mm),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3 * mm),
            ]
        )
    )
    story.append(metric_table)
    story.append(Spacer(1, 5 * mm))
    story.append(
        Paragraph(
            f"<b>Zonage PTZ :</b> {_pdf_escape(ptz_text)}",
            body,
        )
    )
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("Récapitulatif des parcelles", h2))

    summary_data = [
        [
            Paragraph("<b>Référence</b>", small),
            Paragraph("<b>Adresse</b>", small),
            Paragraph("<b>Surf.</b>", small),
            Paragraph("<b>Zone PLU</b>", small),
            Paragraph("<b>PTZ</b>", small),
            Paragraph("<b>Log.</b>", small),
        ]
    ]
    for _, row in rows.iterrows():
        summary_data.append(
            [
                Paragraph(_pdf_escape(row.get("reference")), small),
                Paragraph(_pdf_escape(row.get("adresse")), small),
                Paragraph(_pdf_number(row.get("surface_m2"), 0), small),
                Paragraph(_pdf_escape(row.get("zone_plu")), small),
                Paragraph(_pdf_escape(row.get("zone_ptz")), small),
                Paragraph(_pdf_number(row.get("logements_plu"), 0), small),
            ]
        )

    summary_table = Table(
        summary_data,
        colWidths=[31 * mm, 64 * mm, 19 * mm, 22 * mm, 14 * mm, 17 * mm],
        repeatRows=1,
    )
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), sagec_blue),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.35, line),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 1.5 * mm),
                ("RIGHTPADDING", (0, 0), (-1, -1), 1.5 * mm),
                ("TOPPADDING", (0, 0), (-1, -1), 1.6 * mm),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1.6 * mm),
            ]
        )
    )
    story.append(summary_table)

    # --------------------------------------------------------------
    # Fiches par parcelle
    # --------------------------------------------------------------
    for _, row in rows.iterrows():
        story.append(PageBreak())

        ref = str(row.get("reference") or "")
        address = str(row.get("adresse") or "").strip()
        if not address:
            address = (
                f"Section {row.get('section', '')} n° {row.get('numero', '')} "
                f"- {commune_name}"
            )

        story.append(
            Paragraph(
                f"Parcelle {_pdf_escape(ref)}",
                title_style,
            )
        )
        story.append(
            Paragraph(
                f"<b>{_pdf_escape(address)}</b><br/>"
                f"Section {_pdf_escape(row.get('section'))} - "
                f"N° {_pdf_escape(row.get('numero'))}",
                subtitle_style,
            )
        )

        feature = (feature_map or {}).get(ref)

        key_data = [
            [
                Paragraph("Surface terrain", small),
                Paragraph(f"<b>{_pdf_number(row.get('surface_m2'), 0, ' m²')}</b>", body),
            ],
            [
                Paragraph("Surface en zones constructibles", small),
                Paragraph(f"<b>{_pdf_number(row.get('surface_plu_eligible_m2'), 0, ' m²')}</b>", body),
            ],
            [
                Paragraph("Zone PLU", small),
                Paragraph(f"<b>{_pdf_escape(row.get('zone_plu'))}</b>", body),
            ],
            [
                Paragraph("Zone PTZ", small),
                Paragraph(f"<b>{_pdf_escape(row.get('zone_ptz'))}</b>", body),
            ],
            [
                Paragraph("Terrain", small),
                Paragraph(
                    "<b>Bâti</b>" if bool(row.get("terrain_bati")) else "<b>Nu</b>",
                    body,
                ),
            ],
            [
                Paragraph("Score", small),
                Paragraph(f"<b>{_pdf_number(row.get('score'), 0, ' / 100')}</b>", body),
            ],
        ]
        key_table = Table(key_data, colWidths=[28 * mm, 45 * mm])
        key_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F3F6F8")),
                    ("GRID", (0, 0), (-1, -1), 0.3, line),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 2 * mm),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 2 * mm),
                    ("TOPPADDING", (0, 0), (-1, -1), 2.2 * mm),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2.2 * mm),
                ]
            )
        )

        hero = Table(
            [[ParcelSketch(feature), key_table]],
            colWidths=[84 * mm, 78 * mm],
        )
        hero.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 3 * mm),
                ]
            )
        )
        story.append(hero)
        story.append(Spacer(1, 5 * mm))

        story.append(Paragraph("Potentiel urbanistique", h2))
        urban_data = [
            [
                Paragraph("Emprise PLU", small),
                Paragraph("<b>" + _pdf_number(row.get("emprise_plu_pct"), 1, " %") + "</b>", body),
                Paragraph("Niveaux PLU", small),
                Paragraph("<b>" + _pdf_number(row.get("niveaux_plu"), 0) + "</b>", body),
            ],
            [
                Paragraph("Hauteur PLU", small),
                Paragraph("<b>" + _pdf_number(row.get("hauteur_plu_m"), 1, " m") + "</b>", body),
                Paragraph("Surface brute", small),
                Paragraph("<b>" + _pdf_number(row.get("surface_brute_plu_m2"), 0, " m²") + "</b>", body),
            ],
            [
                Paragraph("SDP estimée", small),
                Paragraph("<b>" + _pdf_number(row.get("sdp_plu_m2"), 0, " m²") + "</b>", body),
                Paragraph("SHAB estimée", small),
                Paragraph("<b>" + _pdf_number(row.get("shab_plu_m2"), 0, " m²") + "</b>", body),
            ],
            [
                Paragraph("Logements estimés", small),
                Paragraph("<b>" + _pdf_number(row.get("logements_plu"), 0) + "</b>", body),
                Paragraph("Couverture emprise + hauteur", small),
                Paragraph("<b>" + _pdf_number(row.get("gabarit_couverture_pct"), 0, " %") + "</b>", body),
            ],
        ]
        urban_table = Table(
            urban_data,
            colWidths=[31 * mm, 48 * mm, 31 * mm, 48 * mm],
        )
        urban_table.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.35, line),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F3F6F8")),
                    ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#F3F6F8")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 2 * mm),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 2 * mm),
                    ("TOPPADDING", (0, 0), (-1, -1), 2 * mm),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2 * mm),
                ]
            )
        )
        story.append(urban_table)
        story.append(Spacer(1, 4 * mm))

        owner = str(row.get("proprietaire_personne_morale") or "").strip()
        owner_type = str(row.get("proprietaire_type") or "").strip()
        siren = str(row.get("siren_proprietaire") or "").strip()

        story.append(Paragraph("Informations de prospection", h2))
        info_rows = [
            [
                Paragraph("Propriétaire société / commune", small),
                Paragraph(
                    _pdf_escape(owner) if owner else "Non identifié comme personne morale",
                    body,
                ),
            ],
            [
                Paragraph("Type propriétaire", small),
                Paragraph(_pdf_escape(owner_type), body),
            ],
            [
                Paragraph("SIREN", small),
                Paragraph(_pdf_escape(siren), body),
            ],
            [
                Paragraph("Bâti existant", small),
                Paragraph(
                    "Oui" if bool(row.get("terrain_bati")) else "Non",
                    body,
                ),
            ],
            [
                Paragraph("Collectif existant", small),
                Paragraph(
                    "Oui" if bool(row.get("collectif_existant")) else "Non",
                    body,
                ),
            ],
        ]
        info_table = Table(info_rows, colWidths=[52 * mm, 106 * mm])
        info_table.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.35, line),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F3F6F8")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 2 * mm),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 2 * mm),
                    ("TOPPADDING", (0, 0), (-1, -1), 2 * mm),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2 * mm),
                ]
            )
        )
        story.append(info_table)
        story.append(Spacer(1, 4 * mm))

        prescription_e = str(row.get("prescription_emprise") or "").strip()
        prescription_h = str(row.get("prescription_hauteur") or "").strip()
        status = str(row.get("gabarit_plu_statut") or "").strip()

        if prescription_e or prescription_h or status:
            story.append(Paragraph("Règles et traçabilité", h2))
            trace_parts = []
            if status:
                trace_parts.append(f"<b>Statut :</b> {_pdf_escape(status)}")
            if prescription_e:
                trace_parts.append(
                    f"<b>Emprise :</b> {_pdf_escape(prescription_e)}"
                )
            if prescription_h:
                trace_parts.append(
                    f"<b>Hauteur :</b> {_pdf_escape(prescription_h)}"
                )
            story.append(
                Paragraph("<br/>".join(trace_parts), small)
            )

        story.append(Spacer(1, 4 * mm))
        story.append(
            Paragraph(
                "<b>Note :</b> ce dossier constitue une présélection foncière. "
                "Les capacités indiquées sont à confirmer par une étude réglementaire et architecturale "
                "complète (retraits, stationnement, OAP, servitudes, risques et autres prescriptions).",
                small,
            )
        )

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)

    buffer.seek(0)
    return buffer




# --------------------------
# Interface
# --------------------------
st.title("🏗️ Prospecteur Foncier — V3.8 — Analyse PLU/PLUi par commune")
st.caption(
    "Cadastre le plus récent + PLU/PLUi en vigueur + analyse précise emprise/hauteur par commune."
)

st.info(
    "Les communes sont chargées depuis l'API administrative officielle. "
    "Le nouveau filtre de zonage PTZ utilise la liste A/B/C officielle du ministère chargé du Logement. "
    "Les parcelles et bâtiments viennent du cadastre réel."
)

with st.sidebar:
    st.header("Territoire")
    region_choice = st.selectbox(
        "Périmètre",
        ["Nouvelle-Aquitaine", "Occitanie", "Tout le Sud-Ouest"],
    )

    if region_choice == "Tout le Sud-Ouest":
        deps = {}
        for d in REGIONS.values():
            deps.update(d)
    else:
        deps = REGIONS[region_choice]

    dep_label_to_code = {f"{code} — {name}": code for code, name in deps.items()}
    dep_label = st.selectbox("Département", list(dep_label_to_code.keys()))
    dep_code = dep_label_to_code[dep_label]

    # Filtre PTZ demandé : placé avant le choix de la commune.
    ptz_choice = st.selectbox(
        "Zonage PTZ",
        ["Tous", "A", "B1", "B2", "C"],
        index=0,
        help=(
            "Zonage A/B/C officiel utilisé notamment pour le PTZ. "
            "Le choix A inclut également A bis, conformément au classement réglementaire."
        ),
    )

    try:
        communes = fetch_communes(dep_code)
    except Exception as exc:
        st.error(f"Impossible de charger les communes : {exc}")
        st.stop()

    ptz_zonage_index = {}
    ptz_source = ""
    try:
        ptz_zonage_index, ptz_source = fetch_ptz_zonage()
        communes = filter_communes_by_ptz(
            communes,
            ptz_choice,
            ptz_zonage_index,
        )
    except Exception as exc:
        if ptz_choice != "Tous":
            st.warning(
                "Le zonage PTZ officiel n'a pas pu être chargé ; "
                "le filtre PTZ est temporairement ignoré. "
                f"Détail : {exc}"
            )

    if not communes:
        st.warning(
            f"Aucune commune du département sélectionné n'est classée en zone PTZ {ptz_choice}."
        )
        st.stop()

    commune_labels = {}
    for c in communes:
        zone_ptz = (
            c.get("zone_ptz")
            or ptz_zonage_index.get(str(c.get("code") or "").zfill(5), {}).get("zone")
            or ""
        )
        suffix = f" — PTZ {zone_ptz}" if zone_ptz else ""
        commune_labels[f"{c['nom']} ({c['code']}){suffix}"] = c

    commune_label = st.selectbox(
        "Commune",
        list(commune_labels.keys()),
        index=0,
    )
    commune = commune_labels[commune_label]
    commune_name = commune["nom"]
    insee = commune["code"]
    code_epci = commune.get("codeEpci")
    zone_ptz_commune = (
        commune.get("zone_ptz")
        or ptz_zonage_index.get(str(insee).zfill(5), {}).get("zone")
        or "Non renseigné"
    )

    st.caption(f"Commune sélectionnée : zone PTZ **{zone_ptz_commune}**")

st.subheader("1. Critères de recherche")

c1, c2 = st.columns(2)
with c1:
    min_log = st.number_input("Logements minimum", min_value=0, value=20, step=1)
with c2:
    max_log = st.number_input("Logements maximum", min_value=0, value=50, step=1)

st.markdown("#### Ratios de calcul de capacité")
r1, r2, r3, r4 = st.columns(4)
with r1:
    ratio_sdp_pct = st.number_input(
        "SDP / surface brute (%)",
        min_value=1.0,
        max_value=100.0,
        value=80.0,
        step=1.0,
    )
with r2:
    ratio_shab_pct = st.number_input(
        "SHAB / SDP (%)",
        min_value=1.0,
        max_value=100.0,
        value=80.0,
        step=1.0,
    )
with r3:
    shab_par_logement = st.number_input(
        "SHAB moyenne / logement (m²)",
        min_value=10.0,
        max_value=250.0,
        value=55.0,
        step=1.0,
    )
with r4:
    floor_height_m = st.number_input(
        "Hauteur / niveau si PLU en mètres",
        min_value=2.4,
        max_value=5.0,
        value=3.0,
        step=0.1,
        help="Utilisé uniquement si la prescription graphique indique une hauteur en mètres sans R+N.",
    )

st.caption(
    "Calcul précis : surface brute = somme des parties constructibles de la parcelle × emprise applicable × niveaux applicables ; "
    f"SDP × {ratio_sdp_pct:.0f} % ; SHAB × {ratio_shab_pct:.0f} % ; "
    f"{shab_par_logement:.0f} m² SHAB/logement."
)

c5, c6 = st.columns(2)
with c5:
    zone_mode = st.selectbox(
        "Zones à analyser",
        ["Zones U uniquement", "Zones U + AUc constructibles"],
        help=(
            "Les zones A et N sont éliminées. Les zones AUs (à urbaniser bloquées) "
            "sont également éliminées automatiquement."
        ),
    )
with c6:
    terrain_mode = st.selectbox(
        "Terrain",
        ["Indifférent", "Terrain nu", "Terrain bâti"],
    )

# Le filtre Habitat strict est supprimé à la demande.
# Les secteurs à vocation économique restent en revanche systématiquement exclus.
habitat_only = False
include_conditionnel = True

st.info(
    f"Analyse V3.8 spécifique à **{commune_name}** : le logiciel récupère le PLU/PLUi en vigueur, "
    "intersecte réellement chaque parcelle du cadastre `latest` avec le règlement graphique, "
    "puis recherche uniquement l'emprise au sol maximale et la hauteur maximale. "
    "Les zones A/N/AUs et les secteurs économiques/industriels sont éliminés avant le calcul."
)

st.caption(
    "Filtre économique actif : les zones dont la vocation dominante est l'activité économique "
    "(activité, industrie, logistique, commerce, artisanat, bureaux, parc d'activités, etc.) "
    "sont éliminées avant même le filtre Habitat."
)

st.caption(
    "Filtre bâti : une parcelle déjà construite reste éligible, sauf si la BDNB identifie "
    "un bâtiment en « Résidentiel collectif » sur cette parcelle."
)

st.warning(
    "La V3.8 se concentre volontairement sur deux règles seulement : **emprise au sol maximale** "
    "et **hauteur maximale**. Les retraits, stationnement, OAP, servitudes et autres règles ne sont "
    "pas intégrés au calcul de capacité de cette version."
)

analyse_button = st.button(
    f"🔎 Analyser le cadastre réel de {commune_name}",
    type="primary",
)

if analyse_button:
    try:
        with st.spinner("Téléchargement des parcelles cadastrales réelles…"):
            parcelles_geojson = fetch_cadastre_layer(insee, "parcelles")
        with st.spinner("Téléchargement des bâtiments cadastraux…"):
            batiments_geojson = fetch_cadastre_layer(insee, "batiments")
        with st.spinner("Recherche des logements collectifs existants via la BDNB…"):
            try:
                bdnb_rows = fetch_bdnb_commune(insee)
                bdnb_parcel_index = build_bdnb_parcel_index(bdnb_rows)
                st.session_state["bdnb_error"] = ""
            except Exception as bdnb_exc:
                bdnb_parcel_index = {}
                st.session_state["bdnb_error"] = str(bdnb_exc)
        with st.spinner("Recherche du document d'urbanisme en vigueur…"):
            commune_geojson = fetch_cadastre_layer(insee, "communes")
            partition, documents = fetch_gpu_document(
                insee,
                commune_geojson,
                code_epci=code_epci,
            )
        with st.spinner(f"Chargement du zonage PLU/PLUi ({partition})…"):
            zones_geojson = fetch_gpu_zones(partition, insee=insee)

        with st.spinner("Chargement des pièces graphiques emprise / hauteur du PLU/PLUi…"):
            graphic_prescriptions = fetch_gpu_graphic_prescriptions(partition, insee=insee)
            graphic_rule_indexes = build_graphic_rule_indexes(graphic_prescriptions)
            nonbuildable_index = build_nonbuildable_graphic_index(
                graphic_prescriptions
            )

        with st.spinner("Pré-interprétation des règles emprise/hauteur propres à la commune…"):
            zone_rule_catalog = build_commune_zone_rule_catalog(
                zones_geojson,
                include_au=(zone_mode == "Zones U + AUc constructibles"),
                floor_height_m=floor_height_m,
            )

        plu_diagnostic = build_plu_commune_diagnostic(
            partition=partition,
            documents=documents,
            zones_geojson=zones_geojson,
            graphic_prescriptions=graphic_prescriptions,
            include_au=(zone_mode == "Zones U + AUc constructibles"),
        )

        if not zones_geojson.get("features"):
            municipality = fetch_gpu_municipality(insee)
            st.error(
                "Aucun zonage PLU/PLUi exploitable n'a été renvoyé pour cette commune. "
                f"Partition résolue : {partition}. Code EPCI : {code_epci or 'non disponible'}."
            )
            st.json(municipality)
            st.stop()

        with st.spinner("Croisement spatial cadastre ↔ PLU ↔ bâtiments…"):
            results, feature_map = analyse_commune(
                commune_name=commune_name,
                insee=insee,
                parcelles_geojson=parcelles_geojson,
                batiments_geojson=batiments_geojson,
                zones_geojson=zones_geojson,
                bdnb_parcel_index=bdnb_parcel_index,
                nonbuildable_index=nonbuildable_index,
                include_au=(zone_mode == "Zones U + AUc constructibles"),
                habitat_only=habitat_only,
                include_conditionnel=include_conditionnel,
                ratio_sdp_pct=ratio_sdp_pct,
                ratio_shab_pct=ratio_shab_pct,
                shab_par_logement=shab_par_logement,
            )

        with st.spinner("Comparaison précise cadastre ↔ zones PLU ↔ emprise 38 ↔ hauteur 39…"):
            results = enrich_with_precise_commune_rules(
                results,
                feature_map=feature_map,
                rule_indexes=graphic_rule_indexes,
                zone_rule_catalog=zone_rule_catalog,
                ratio_sdp_pct=ratio_sdp_pct,
                ratio_shab_pct=ratio_shab_pct,
                shab_par_logement=shab_par_logement,
                floor_height_m=floor_height_m,
            )

            if feature_map and not results.empty:
                rule_by_ref = results.set_index("reference").to_dict("index")
                for ref, feat in feature_map.items():
                    row = rule_by_ref.get(ref)
                    if not row:
                        continue
                    props = feat.setdefault("properties", {})
                    props["emprise_plu"] = row.get("emprise_plu_pct")
                    props["niveaux_plu"] = row.get("niveaux_plu")
                    props["hauteur_plu"] = row.get("hauteur_plu_m")
                    props["surface_brute_plu"] = row.get("surface_brute_plu_m2")
                    props["sdp_plu"] = row.get("sdp_plu_m2")
                    props["shab_plu"] = row.get("shab_plu_m2")
                    props["logements_plu"] = row.get("logements_plu")
                    props["statut_gabarit"] = row.get("gabarit_plu_statut")
                    props["couverture_regles"] = row.get("gabarit_couverture_pct")
                    props["surface_constructible_plu"] = row.get("surface_plu_eligible_m2")

        with st.spinner("Recherche des propriétaires personnes morales (sociétés / communes)…"):
            try:
                pm_rows = fetch_personnes_morales_commune(insee)
                pm_index = build_personnes_morales_index(insee, pm_rows)
                results = enrich_with_personnes_morales(results, insee, pm_index)

                if feature_map and not results.empty:
                    owner_by_ref = results.set_index("reference")[
                        ["proprietaire_personne_morale", "proprietaire_type"]
                    ].to_dict("index")
                    for ref, feat in feature_map.items():
                        owner = owner_by_ref.get(ref, {})
                        feat.setdefault("properties", {})["proprietaire"] = (
                            owner.get("proprietaire_personne_morale") or "Non identifié (personne morale)"
                        )
                        feat["properties"]["type_proprietaire"] = owner.get("proprietaire_type") or ""

                st.session_state["owner_error"] = ""
            except Exception as owner_exc:
                results = enrich_with_personnes_morales(results, insee, {})
                if feature_map:
                    for _, feat in feature_map.items():
                        feat.setdefault("properties", {})["proprietaire"] = "Non disponible"
                        feat["properties"]["type_proprietaire"] = ""
                st.session_state["owner_error"] = str(owner_exc)

        st.session_state["analysis_results"] = results
        st.session_state["feature_map"] = feature_map
        st.session_state["analysis_insee"] = insee
        st.session_state["analysis_commune"] = commune_name
        st.session_state["analysis_partition"] = partition
        st.session_state["plu_diagnostic"] = plu_diagnostic

    except Exception as exc:
        st.error(
            "L'analyse n'a pas pu être terminée. Les API publiques peuvent parfois être temporairement "
            "indisponibles ou une commune peut ne pas avoir de zonage GPU exploitable."
        )
        st.exception(exc)

results = st.session_state.get("analysis_results")
if results is not None and st.session_state.get("analysis_insee") == insee:
    if results.empty or "logements_estimes" not in results.columns:
        st.subheader("2. Résultats sur le cadastre réel")
        st.warning(
            "Aucune parcelle n'a été retenue pour cette commune avec les règles actuelles "
            "(zones U/AUc, exclusion des secteurs économiques et du collectif existant)."
        )
        st.stop()

    results = results.copy()
    results["logements_plu"] = pd.to_numeric(results["logements_plu"], errors="coerce")

    calculable = results["logements_plu"].notna()

    filtered = results[
        calculable
        & (results["logements_plu"] >= min_log)
        & (results["logements_plu"] <= max_log)
    ].copy()

    unresolved = results[~calculable].copy()

    cleaned = []
    for frame in [filtered, unresolved]:
        if "economic_vocation" in frame.columns:
            frame = frame[frame["economic_vocation"] == False].copy()
        if "collectif_existant" in frame.columns:
            frame = frame[frame["collectif_existant"] == False].copy()

        if terrain_mode == "Terrain nu":
            frame = frame[frame["terrain_bati"] == False]
        elif terrain_mode == "Terrain bâti":
            frame = frame[frame["terrain_bati"] == True]

        frame = frame.sort_values(
            ["score", "surface_m2"],
            ascending=[False, False],
        ).reset_index(drop=True)
        cleaned.append(frame)

    filtered, unresolved = cleaned

    st.subheader("2. Résultats sur le cadastre réel")

    if filtered.empty and not unresolved.empty:
        st.warning(
            "Les parcelles sont bien situées en zones constructibles non industrielles, "
            "mais l'emprise et/ou la hauteur n'ont pas pu être chiffrées avec suffisamment de couverture. "
            "Elles restent visibles dans « Emprise/hauteur à vérifier »."
        )
    elif filtered.empty and unresolved.empty:
        st.warning("Aucune parcelle ne correspond aux critères actuels.")

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Parcelles dans la cible logements", len(filtered))
    m2.metric("Emprise/hauteur à vérifier", len(unresolved))
    m3.metric("Parcelles candidates analysées", len(results))
    m4.metric("Document urbanisme", st.session_state.get("analysis_partition", "—"))

    diag = st.session_state.get("plu_diagnostic", {})
    if diag:
        date_doc = diag.get("date_document") or "date non renseignée"
        st.success(
            f"Document analysé pour {commune_name} : {diag.get('typedoc', 'PLU/PLUi')} "
            f"— partition {diag.get('partition', '—')} — validation/approbation {date_doc}. "
            f"{diag.get('zones_retenues', 0)} zone(s) constructible(s) retenue(s), "
            f"{diag.get('zones_non_constructibles_exclues', 0)} zone(s) non constructible(s) exclue(s), "
            f"{diag.get('zones_economiques_exclues', 0)} zone(s) économique(s)/industrielle(s) exclue(s)."
        )
        st.caption(
            f"Pièces graphiques détectées : {diag.get('prescriptions_emprise', 0)} prescription(s) "
            f"d'emprise 38, {diag.get('prescriptions_hauteur', 0)} prescription(s) de hauteur 39 "
            f"et {diag.get('prescriptions_inconstructibles', 0)} secteur(s) inconstructible(s) 02-01. "
            "Cadastre utilisé : flux Etalab `latest`."
        )

    if st.session_state.get("analysis_partition", "").startswith("DU_") and code_epci:
        if st.session_state.get("analysis_partition") == f"DU_{code_epci}":
            st.caption(
                "Document détecté comme PLUi intercommunal via le code EPCI "
                f"{code_epci}. Le zonage affiché est filtré sur la commune {insee}."
            )

    if st.session_state.get("bdnb_error"):
        st.warning(
            "La BDNB n'a pas répondu pendant cette analyse. Dans ce cas, le logiciel ne peut pas "
            "garantir l'exclusion des résidences collectives. Détail : "
            + st.session_state["bdnb_error"]
        )

    if st.session_state.get("owner_error"):
        st.warning(
            "La recherche des propriétaires personnes morales n'a pas pu être chargée. "
            "Les parcelles restent analysées normalement, mais le nom du propriétaire société/commune "
            "peut être vide. Détail : " + st.session_state["owner_error"]
        )

    map_source = filtered if not filtered.empty else unresolved

    if map_source.empty:
        st.warning("Aucune parcelle ne correspond aux critères actuels.")
    else:
        fmap = st.session_state.get("feature_map", {})
        map_features = []
        for ref in map_source["reference"].head(2000):
            feat = fmap.get(ref)
            if feat:
                map_features.append(feat)

        if map_features:
            fc = {"type": "FeatureCollection", "features": map_features}
            center_lat = float(map_source["latitude"].mean())
            center_lon = float(map_source["longitude"].mean())

            layer = pdk.Layer(
                "GeoJsonLayer",
                fc,
                pickable=True,
                stroked=True,
                filled=True,
                opacity=0.25,
                get_line_width=1,
            )
            deck = pdk.Deck(
                layers=[layer],
                initial_view_state=pdk.ViewState(
                    latitude=center_lat,
                    longitude=center_lon,
                    zoom=12,
                    pitch=0,
                ),
                tooltip={
                    "html": (
                        "<b>{reference}</b><br/>"
                        "Surface terrain : {surface_m2} m²<br/>"
                        "Part PLU constructible : {surface_constructible_plu} m²<br/>"
                        "Emprise PLU : {emprise_plu} %<br/>"
                        "Niveaux PLU : {niveaux_plu}<br/>"
                        "Hauteur PLU : {hauteur_plu} m<br/>"
                        "Surface brute PLU : {surface_brute_plu} m²<br/>"
                        "SDP PLU : {sdp_plu} m²<br/>"
                        "SHAB PLU : {shab_plu} m²<br/>"
                        "Habitat : {habitat}<br/>"
                        "Confiance habitat : {confiance_habitat} %<br/>"
                        "Collectif existant : {collectif}<br/>"
                        "Propriétaire PM : {proprietaire}<br/>"
                        "Type propriétaire : {type_proprietaire}<br/>"
                        "Zone : {typezone} {zone}<br/>"
                        "Potentiel PLU : {logements_plu} logements<br/>"
                        "Gabarit : {statut_gabarit}<br/>"
                        "Couverture emprise+hauteur : {couverture_regles} %"
                    )
                },
            )
            st.pydeck_chart(deck, use_container_width=True)

        st.caption(
            "La carte affiche les parcelles cadastrales croisées avec les prescriptions graphiques du PLU."
        )

        if not unresolved.empty:
            with st.expander(
                f"Gabarit PLU à vérifier ({len(unresolved)} parcelles)",
                expanded=filtered.empty,
            ):
                cols = [
                    "reference",
                    "section",
                    "numero",
                    "surface_m2",
                    "zone_plu",
                    "emprise_plu_pct",
                    "niveaux_plu",
                    "hauteur_plu_m",
                    "gabarit_plu_statut",
                    "prescription_emprise",
                    "prescription_hauteur",
                    "prescription_url",
                ]
                st.dataframe(
                    unresolved[[c for c in cols if c in unresolved.columns]],
                    use_container_width=True,
                    hide_index=True,
                )
                st.caption(
                    "Ces parcelles ne sont pas supprimées : le règlement graphique ne fournit simplement "
                    "pas une valeur chiffrée exploitable sur toute la parcelle."
                )

        st.subheader("3. Sélection des parcelles")
        if filtered.empty:
            st.info(
                "Aucune parcelle n'est encore sélectionnable par nombre de logements calculé. "
                "Consulte la liste « Gabarit PLU à vérifier »."
            )
            st.stop()

        table = filtered[
            [
                "selection",
                "reference",
                "section",
                "numero",
                "proprietaire_personne_morale",
                "proprietaire_type",
                "forme_juridique_proprietaire",
                "siren_proprietaire",
                "surface_m2",
                "surface_plu_eligible_m2",
                "part_constructible_plu_pct",
                "zones_intersectees",
                "surface_inconstructible_graphique_m2",
                "prescriptions_inconstructibles",
                "emprise_plu_pct",
                "emprise_graphique_couverture_pct",
                "niveaux_plu",
                "hauteur_plu_m",
                "surface_brute_plu_m2",
                "sdp_plu_m2",
                "shab_plu_m2",
                "terrain_bati",
                "collectif_existant",
                "usage_bdnb",
                "nb_batiments",
                "zone_type",
                "zone_plu",
                "zone_ptz",
                "habitat_statut",
                "habitat_preuve",
                "habitat_confiance",
                "formdomi",
                "economic_vocation",
                "classe_zone",
                "logements_plu",
                "gabarit_plu_statut",
                "gabarit_plu_confiance",
                "gabarit_couverture_pct",
                "regle_emprise_source",
                "regle_hauteur_source",
                "regle_emprise_extrait",
                "regle_hauteur_extrait",
                "prescription_emprise",
                "prescription_hauteur",
                "prescription_url",
                "score",
                "reglement_url",
                "latitude",
                "longitude",
                "adresse",
            ]
        ].copy()

        edited = st.data_editor(
            table,
            hide_index=True,
            use_container_width=True,
            column_config={
                "selection": st.column_config.CheckboxColumn("Sélectionner"),
                "reference": st.column_config.TextColumn("Référence"),
                "section": st.column_config.TextColumn("Section"),
                "numero": st.column_config.TextColumn("N°"),
                "proprietaire_personne_morale": st.column_config.TextColumn("Propriétaire société / commune"),
                "proprietaire_type": st.column_config.TextColumn("Type propriétaire"),
                "forme_juridique_proprietaire": st.column_config.TextColumn("Forme juridique"),
                "siren_proprietaire": st.column_config.TextColumn("SIREN"),
                "surface_m2": st.column_config.NumberColumn("Surface terrain m²"),
                "surface_plu_eligible_m2": st.column_config.NumberColumn("Surface en zones constructibles m²"),
                "part_constructible_plu_pct": st.column_config.NumberColumn("Part constructible PLU %", format="%.1f %%"),
                "zones_intersectees": st.column_config.TextColumn("Zones PLU intersectées"),
                "surface_inconstructible_graphique_m2": st.column_config.NumberColumn("Surface exclue 02-01 m²"),
                "prescriptions_inconstructibles": st.column_config.TextColumn("Secteurs inconstructibles"),
                "emprise_plu_pct": st.column_config.NumberColumn("Emprise PLU %", format="%.1f %%"),
                "emprise_graphique_couverture_pct": st.column_config.NumberColumn("Couverture emprise %", format="%.1f %%"),
                "niveaux_plu": st.column_config.NumberColumn("Niveaux PLU"),
                "hauteur_plu_m": st.column_config.NumberColumn("Hauteur PLU m", format="%.1f"),
                "surface_brute_plu_m2": st.column_config.NumberColumn("Surface brute PLU m²"),
                "sdp_plu_m2": st.column_config.NumberColumn("SDP PLU m²"),
                "shab_plu_m2": st.column_config.NumberColumn("SHAB PLU m²"),
                "terrain_bati": st.column_config.CheckboxColumn("Bâti"),
                "collectif_existant": st.column_config.CheckboxColumn("Collectif existant"),
                "usage_bdnb": st.column_config.TextColumn("Usage BDNB"),
                "nb_batiments": st.column_config.NumberColumn("Nb bâtiments"),
                "zone_type": st.column_config.TextColumn("Type zone"),
                "zone_plu": st.column_config.TextColumn("Zone PLU"),
                "zone_ptz": st.column_config.TextColumn("Zone PTZ"),
                "habitat_statut": st.column_config.TextColumn("Statut habitat"),
                "habitat_preuve": st.column_config.TextColumn("Preuve habitat"),
                "habitat_confiance": st.column_config.ProgressColumn(
                    "Confiance habitat", min_value=0, max_value=100
                ),
                "formdomi": st.column_config.TextColumn("Forme dominante CNIG"),
                "economic_vocation": st.column_config.CheckboxColumn("Vocation économique"),
                "classe_zone": st.column_config.TextColumn("Qualification"),
                "logements_plu": st.column_config.NumberColumn("Logements PLU"),
                "gabarit_plu_statut": st.column_config.TextColumn("Statut gabarit"),
                "gabarit_plu_confiance": st.column_config.ProgressColumn("Confiance gabarit", min_value=0, max_value=100),
                "gabarit_couverture_pct": st.column_config.ProgressColumn("Couverture règles", min_value=0, max_value=100),
                "regle_emprise_source": st.column_config.TextColumn("Source emprise"),
                "regle_hauteur_source": st.column_config.TextColumn("Source hauteur"),
                "regle_emprise_extrait": st.column_config.TextColumn("Extrait règle emprise"),
                "regle_hauteur_extrait": st.column_config.TextColumn("Extrait règle hauteur"),
                "prescription_emprise": st.column_config.TextColumn("Prescription emprise"),
                "prescription_hauteur": st.column_config.TextColumn("Prescription hauteur"),
                "prescription_url": st.column_config.LinkColumn("Prescription"),
                "score": st.column_config.ProgressColumn("Score", min_value=0, max_value=100),
                "reglement_url": st.column_config.LinkColumn("Règlement"),
                "latitude": None,
                "longitude": None,
                "adresse": st.column_config.TextColumn("Adresse"),
            },
            disabled=[
                "reference",
                "section",
                "numero",
                "proprietaire_personne_morale",
                "proprietaire_type",
                "forme_juridique_proprietaire",
                "siren_proprietaire",
                "surface_m2",
                "surface_plu_eligible_m2",
                "part_constructible_plu_pct",
                "zones_intersectees",
                "surface_inconstructible_graphique_m2",
                "prescriptions_inconstructibles",
                "emprise_plu_pct",
                "emprise_graphique_couverture_pct",
                "niveaux_plu",
                "hauteur_plu_m",
                "surface_brute_plu_m2",
                "sdp_plu_m2",
                "shab_plu_m2",
                "terrain_bati",
                "collectif_existant",
                "usage_bdnb",
                "nb_batiments",
                "zone_type",
                "zone_plu",
                "habitat_statut",
                "habitat_preuve",
                "habitat_confiance",
                "formdomi",
                "economic_vocation",
                "classe_zone",
                "logements_plu",
                "gabarit_plu_statut",
                "gabarit_plu_confiance",
                "gabarit_couverture_pct",
                "regle_emprise_source",
                "regle_hauteur_source",
                "regle_emprise_extrait",
                "regle_hauteur_extrait",
                "prescription_emprise",
                "prescription_hauteur",
                "prescription_url",
                "score",
                "reglement_url",
                "latitude",
                "longitude",
                "adresse",
            ],
            key="parcel_editor",
        )

        selected = edited[edited["selection"] == True].copy()
        st.write(f"**{len(selected)} parcelle(s) sélectionnée(s)**")

        if not selected.empty:
            if st.button("📍 Rechercher l'adresse des parcelles sélectionnées"):
                progress = st.progress(0)
                addresses = {}
                for i, (_, row) in enumerate(selected.iterrows(), start=1):
                    existing_addr = str(row.get("adresse") or "").strip()
                    if existing_addr:
                        addr = existing_addr
                    else:
                        try:
                            addr = reverse_geocode(
                                float(row["latitude"]),
                                float(row["longitude"]),
                                insee,
                            )
                        except Exception:
                            addr = ""
                    addresses[row["reference"]] = addr
                    progress.progress(i / len(selected))
                    time.sleep(0.03)

                # Conserver les adresses dans les résultats de session.
                base_results = st.session_state["analysis_results"].copy()
                for ref, addr in addresses.items():
                    base_results.loc[
                        base_results["reference"] == ref, "adresse"
                    ] = addr
                st.session_state["analysis_results"] = base_results
                st.success(
                    "Adresses recherchées. La page va réutiliser ces adresses pour les courriers."
                )
                st.rerun()

            st.subheader("4. Dossier PDF de la sélection")

            # Toujours reprendre les données les plus récentes (notamment les adresses).
            latest_pdf = st.session_state["analysis_results"].set_index("reference")
            pdf_rows = []
            for _, sel_row in selected.iterrows():
                ref = sel_row["reference"]
                if ref in latest_pdf.index:
                    current = latest_pdf.loc[ref]
                    if isinstance(current, pd.DataFrame):
                        current = current.iloc[0]
                    pdf_rows.append(current.to_dict())
                else:
                    pdf_rows.append(sel_row.to_dict())

            pdf_selection_df = pd.DataFrame(pdf_rows)

            try:
                pdf_buffer = generate_selected_parcels_pdf(
                    pdf_selection_df,
                    commune_name=commune_name,
                    insee=insee,
                    feature_map=st.session_state.get("feature_map", {}),
                )
                st.download_button(
                    "📄 Télécharger le dossier PDF des parcelles sélectionnées",
                    data=pdf_buffer.getvalue(),
                    file_name=(
                        f"selection_fonciere_{commune_name}_{date.today().isoformat()}.pdf"
                    ),
                    mime="application/pdf",
                )
                st.caption(
                    "Le PDF contient une synthèse générale puis une fiche par parcelle "
                    "avec schéma cadastral, adresse, propriétaire personne morale, zonages "
                    "et potentiel urbanistique."
                )
            except Exception as pdf_exc:
                st.warning(
                    "Le dossier PDF n'a pas pu être généré. "
                    f"Détail : {pdf_exc}"
                )

            st.subheader("5. Courriers de prospection")
            l1, l2 = st.columns(2)
            with l1:
                signataire = st.text_input("Signataire", "Nicolas PEDROT")
                fonction = st.text_input("Fonction", "Directeur")
            with l2:
                email = st.text_input("E-mail", "Nicolas.pedrot@sagec.fr")
                ville_signature = st.text_input("Ville de signature", "Anglet")

            if st.button("✉️ Générer les lettres sélectionnées"):
                # Récupérer les adresses éventuellement trouvées dans le DataFrame principal.
                latest = st.session_state["analysis_results"].set_index("reference")
                zip_buffer = io.BytesIO()
                generated = 0
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as z:
                    for _, row in selected.iterrows():
                        ref = row["reference"]
                        if ref in latest.index:
                            current = latest.loc[ref].to_dict()
                        else:
                            current = row.to_dict()

                        letter = generate_letter(
                            current,
                            signataire,
                            fonction,
                            email,
                            ville_signature,
                        )
                        safe_ref = re.sub(r"[^A-Za-z0-9_-]+", "_", str(ref))
                        z.writestr(
                            f"Lettre_{commune_name}_{safe_ref}.docx",
                            letter.getvalue(),
                        )
                        generated += 1

                zip_buffer.seek(0)
                st.download_button(
                    f"Télécharger {generated} lettre(s)",
                    data=zip_buffer.getvalue(),
                    file_name=f"courriers_{commune_name}_{date.today().isoformat()}.zip",
                    mime="application/zip",
                )

        st.download_button(
            "⬇️ Exporter les résultats en CSV",
            data=filtered.to_csv(index=False).encode("utf-8-sig"),
            file_name=f"prospection_{commune_name}_{insee}.csv",
            mime="text/csv",
        )

with st.expander("Ce que la V2 analyse réellement"):
    st.markdown(
        """
- **Communes** : chargées dynamiquement pour la Nouvelle-Aquitaine et l'Occitanie.
- **Cadastre** : parcelles et bâtiments réels du dernier millésime Etalab.
- **PLU / PLUi** : zonages du Géoportail de l'Urbanisme via l'API Carto IGN.
- **Exclusion économique prioritaire** : `DESTDOMI=02`, `FORMDOMI=0200/0201/0202/0203` et libellés explicites d'activités économiques sont éliminés.
- **Règlement graphique** : croisement spatial des parcelles avec `PRESCRIPTION_SURF`, `PRESCRIPTION_LIN` et `PRESCRIPTION_PCT`.
- **Emprise maximale** : prescription CNIG 38-02.
- **Hauteur maximale** : prescription CNIG 39-02 (et variantes localisées quand elles sont publiées).
- **Prudence** : une prescription ne couvrant qu'une petite partie de la parcelle n'est pas généralisée à toute la parcelle.
- **Propriétaires personnes morales** : tentative de rapprochement avec le fichier open data DGFiP des parcelles détenues par des personnes morales ; affichage du nom, de la forme juridique et du SIREN lorsqu'ils sont disponibles.
- **Compatibilité anciens PLU** : lecture de `DESTDOMI` (01 habitat ; 03 mixte habitat/activité) puis analyse des libellés explicites.
- **Exclusions de zonage** : A, N et AUs sont écartées ; U et AUc sont analysées selon la destination logement.
- **Mode haute précision** : une zone ambiguë sans preuve explicite d'habitat est exclue.
- **Terrain nu / bâti** : déterminé par croisement spatial avec les bâtiments cadastraux.
- **Collectif existant** : les parcelles identifiées par la BDNB comme déjà occupées par du « Résidentiel collectif » sont exclues.
- **Autres parcelles bâties** : elles restent éligibles, quelle que soit la taille ou l'emprise du bâtiment.
- **Capacité logements** : SDP = surface brute × ratio SDP ; SHAB = SDP × ratio SHAB ; logements = SHAB ÷ ratio SHAB/logement.
- **Adresse** : recherchée pour les parcelles sélectionnées via le géocodage inverse de la Géoplateforme.
- **Courrier** : généré à partir du nouveau modèle SAGEC amélioré fourni.
- **Dossier PDF** : synthèse de la sélection + une fiche détaillée par parcelle avec schéma cadastral.

### Limite actuelle
Le filtre Habitat exploite désormais les destinations structurées du PLU lorsqu'elles sont disponibles.
Il ne remplace toutefois pas une faisabilité complète : le nombre de logements reste une estimation de
**présélection**. Une prochaine version pourra extraire automatiquement du règlement : emprise au sol,
hauteur, retraits, pleine terre, stationnement, OAP, prescriptions et servitudes.
        """
    )
